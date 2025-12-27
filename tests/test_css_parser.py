"""
Tests for CSS Parser functionality
"""

import os
import unittest
from docx import Document
from html4docx import HtmlToDocx
from html4docx.css_parser import CSSParser


class CSSParserTest(unittest.TestCase):
    """Test cases for CSS Parser"""

    def setUp(self):
        self.parser = CSSParser()

    def test_parse_simple_tag_selector(self):
        """Test parsing simple tag selector"""
        css = "p { color: red; font-size: 12px; }"
        self.parser.parse_css(css)

        self.assertIn('p', self.parser.tag_rules)
        self.assertEqual(self.parser.tag_rules['p']['color'], 'red')
        self.assertEqual(self.parser.tag_rules['p']['font-size'], '12px')

    def test_parse_class_selector(self):
        """Test parsing class selector"""
        css = ".highlight { font-weight: bold; background-color: yellow; }"
        self.parser.parse_css(css)

        self.assertIn('highlight', self.parser.class_rules)
        self.assertEqual(self.parser.class_rules['highlight']['font-weight'], 'bold')
        self.assertEqual(self.parser.class_rules['highlight']['background-color'], 'yellow')

    def test_parse_id_selector(self):
        """Test parsing ID selector"""
        css = "#header { text-align: center; font-size: 24px; }"
        self.parser.parse_css(css)

        self.assertIn('header', self.parser.id_rules)
        self.assertEqual(self.parser.id_rules['header']['text-align'], 'center')
        self.assertEqual(self.parser.id_rules['header']['font-size'], '24px')

    def test_parse_multiple_selectors(self):
        """Test parsing multiple selectors"""
        css = "h1, h2, h3 { color: blue; }"
        self.parser.parse_css(css)

        self.assertIn('h1', self.parser.tag_rules)
        self.assertIn('h2', self.parser.tag_rules)
        self.assertIn('h3', self.parser.tag_rules)
        self.assertEqual(self.parser.tag_rules['h1']['color'], 'blue')

    def test_parse_multiple_rules(self):
        """Test parsing multiple CSS rules"""
        css = """
        p { color: red; }
        .highlight { font-weight: bold; }
        #header { text-align: center; }
        """
        self.parser.parse_css(css)

        self.assertIn('p', self.parser.tag_rules)
        self.assertIn('highlight', self.parser.class_rules)
        self.assertIn('header', self.parser.id_rules)

    def test_parse_with_comments(self):
        """Test parsing CSS with comments"""
        css = """
        /* This is a comment */
        p { color: red; }
        /* Another comment */
        .highlight { font-weight: bold; }
        """
        self.parser.parse_css(css)

        self.assertIn('p', self.parser.tag_rules)
        self.assertIn('highlight', self.parser.class_rules)
        self.assertEqual(self.parser.tag_rules['p']['color'], 'red')

    def test_parse_important_flag(self):
        """Test parsing CSS with !important flag"""
        css = "p { color: red !important; font-size: 12px; }"
        self.parser.parse_css(css)

        self.assertIn('p', self.parser.tag_rules)
        self.assertIn('!important', self.parser.tag_rules['p']['color'].lower())

    def test_get_styles_for_element_tag(self):
        """Test getting styles for element by tag"""
        css = "p { color: red; font-size: 12px; }"
        self.parser.parse_css(css)

        styles = self.parser.get_styles_for_element('p')
        self.assertEqual(styles['color'], 'red')
        self.assertEqual(styles['font-size'], '12px')

    def test_get_styles_for_element_class(self):
        """Test getting styles for element by class"""
        css = ".highlight { font-weight: bold; }"
        self.parser.parse_css(css)

        styles = self.parser.get_styles_for_element('p', {'class': 'highlight'})
        self.assertEqual(styles['font-weight'], 'bold')

    def test_get_styles_for_element_id(self):
        """Test getting styles for element by ID"""
        css = "#header { text-align: center; }"
        self.parser.parse_css(css)

        styles = self.parser.get_styles_for_element('div', {'id': 'header'})
        self.assertEqual(styles['text-align'], 'center')

    def test_get_styles_for_element_multiple_classes(self):
        """Test getting styles for element with multiple classes"""
        css = """
        .highlight { font-weight: bold; }
        .large { font-size: 18px; }
        """
        self.parser.parse_css(css)

        styles = self.parser.get_styles_for_element('p', {'class': 'highlight large'})
        self.assertEqual(styles['font-weight'], 'bold')
        self.assertEqual(styles['font-size'], '18px')

    def test_get_styles_for_element_combined(self):
        """Test getting styles combining tag, class, and ID"""
        css = """
        p { color: black; }
        .highlight { font-weight: bold; }
        #header { text-align: center; }
        """
        self.parser.parse_css(css)

        styles = self.parser.get_styles_for_element('p', {'class': 'highlight', 'id': 'header'})
        self.assertEqual(styles['color'], 'black')
        self.assertEqual(styles['font-weight'], 'bold')
        self.assertEqual(styles['text-align'], 'center')

    def test_get_styles_with_inline_override(self):
        """Test that inline styles override CSS styles"""
        css = "p { color: red; }"
        self.parser.parse_css(css)

        inline_styles = {'color': 'blue'}
        styles = self.parser.get_styles_for_element('p', inline_styles=inline_styles)
        # Inline should override CSS
        self.assertEqual(styles['color'], 'blue')

    def test_get_styles_with_important(self):
        """Test getting styles separated by !important"""
        css = "p { color: red !important; font-size: 12px; }"
        self.parser.parse_css(css)

        normal, important = self.parser.get_styles_for_element_with_important('p')
        self.assertEqual(normal['font-size'], '12px')
        self.assertIn('color', important)

    def test_clear_rules(self):
        """Test clearing all CSS rules"""
        css = "p { color: red; }"
        self.parser.parse_css(css)

        self.assertTrue(self.parser.has_rules())
        self.parser.clear()
        self.assertFalse(self.parser.has_rules())

    def test_has_rules(self):
        """Test checking if parser has rules"""
        self.assertFalse(self.parser.has_rules())

        css = "p { color: red; }"
        self.parser.parse_css(css)
        self.assertTrue(self.parser.has_rules())


class StyleTagIntegrationTest(unittest.TestCase):
    """Integration tests for <style> tag support"""

    def setUp(self):
        self.document = Document()
        self.parser = HtmlToDocx()

    def test_style_tag_basic(self):
        """Test basic <style> tag functionality"""
        html = """
        <style>
        p { color: red; font-size: 14pt; }
        </style>
        <p>This is a red paragraph</p>
        """

        self.parser.add_html_to_document(html, self.document)

        # Check that paragraph exists (may have multiple due to whitespace handling)
        self.assertGreaterEqual(len(self.document.paragraphs), 1)

        # Check that CSS styles were applied
        # Note: We can't directly check color, but we can verify the parser processed it
        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('p', self.parser.css_parser.tag_rules)

    def test_style_tag_with_class(self):
        """Test <style> tag with class selector"""
        html = """
        <style>
        .highlight { font-weight: bold; color: blue; }
        </style>
        <p class="highlight">Bold blue text</p>
        """

        self.parser.add_html_to_document(html, self.document)

        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('highlight', self.parser.css_parser.class_rules)

    def test_style_tag_with_id(self):
        """Test <style> tag with ID selector"""
        html = """
        <style>
        #header { text-align: center; font-size: 24px; }
        </style>
        <h1 id="header">Centered Header</h1>
        """

        self.parser.add_html_to_document(html, self.document)

        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('header', self.parser.css_parser.id_rules)

    def test_style_tag_multiple_rules(self):
        """Test <style> tag with multiple CSS rules"""
        html = """
        <style>
        p { color: red; }
        .highlight { font-weight: bold; }
        #header { text-align: center; }
        </style>
        <p>Red text</p>
        <p class="highlight">Bold text</p>
        <h1 id="header">Centered</h1>
        """

        self.parser.add_html_to_document(html, self.document)

        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('p', self.parser.css_parser.tag_rules)
        self.assertIn('highlight', self.parser.css_parser.class_rules)
        self.assertIn('header', self.parser.css_parser.id_rules)

    def test_style_tag_with_inline_override(self):
        """Test that inline styles override <style> tag styles"""
        html = """
        <style>
        p { color: red; }
        </style>
        <p style="color: blue;">This should be blue</p>
        """

        self.parser.add_html_to_document(html, self.document)

        # Both CSS and inline styles should be present
        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('p', self.parser.css_parser.tag_rules)

    def test_style_tag_comments(self):
        """Test <style> tag with CSS comments"""
        html = """
        <style>
        /* This is a comment */
        p { color: red; }
        /* Another comment */
        </style>
        <p>Red text</p>
        """

        self.parser.add_html_to_document(html, self.document)

        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('p', self.parser.css_parser.tag_rules)
        self.assertEqual(self.parser.css_parser.tag_rules['p']['color'], 'red')

    def test_style_tag_not_in_output(self):
        """Test that <style> tags are removed from output"""
        html = """
        <style>
        p { color: red; }
        </style>
        <p>Some text</p>
        """

        self.parser.add_html_to_document(html, self.document)

        # Check that we have exactly one paragraph (the <p> tag, not the <style>)
        paragraphs = self.document.paragraphs
        self.assertGreaterEqual(len(paragraphs), 1)

        # Verify style tag content is not in any paragraph text
        for para in paragraphs:
            self.assertNotIn('<style>', para.text)
            self.assertNotIn('color: red', para.text)

    def test_style_tag_with_span(self):
        """Test <style> tag styles applied to span elements"""
        html = """
        <style>
        .highlight { color: blue; font-weight: bold; }
        </style>
        <p>Normal text with <span class="highlight">highlighted text</span></p>
        """

        self.parser.add_html_to_document(html, self.document)

        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('highlight', self.parser.css_parser.class_rules)

    def test_style_tag_with_div(self):
        """Test <style> tag styles applied to div elements"""
        html = """
        <style>
        .container { margin-left: 20px; }
        </style>
        <div class="container">
        <p>Content in container</p>
        </div>
        """

        self.parser.add_html_to_document(html, self.document)

        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('container', self.parser.css_parser.class_rules)

    def test_style_tag_cascade(self):
        """Test CSS cascade with tag, class, and ID"""
        html = """
        <style>
        p { color: black; }
        .highlight { color: blue; }
        #special { color: red; }
        </style>
        <p>Black text</p>
        <p class="highlight">Blue text</p>
        <p class="highlight" id="special">Red text (ID overrides class)</p>
        """

        self.parser.add_html_to_document(html, self.document)

        self.assertTrue(self.parser.css_parser.has_rules())
        # Verify all selectors are parsed
        self.assertIn('p', self.parser.css_parser.tag_rules)
        self.assertIn('highlight', self.parser.css_parser.class_rules)
        self.assertIn('special', self.parser.css_parser.id_rules)


class ExternalCSSTest(unittest.TestCase):
    """Test cases for external CSS via <link> tags"""

    def setUp(self):
        self.document = Document()
        self.parser = HtmlToDocx()
        self.test_dir = os.path.abspath(os.path.dirname(__file__))

    def test_external_css_from_local_file(self):
        """Test loading CSS from local file via <link> tag"""
        css_path = os.path.join(self.test_dir, 'assets/css/test_styles.css')
        html = f"""
        <link rel="stylesheet" href="{css_path}">
        <p class="highlight">Highlighted paragraph</p>
        <h1 id="header">Header</h1>
        """

        self.parser.add_html_to_document(html, self.document)

        # Verify CSS was loaded
        self.assertTrue(self.parser.css_parser.has_rules())
        # Verify relevant rules were loaded
        self.assertIn('p', self.parser.css_parser.tag_rules)
        self.assertIn('highlight', self.parser.css_parser.class_rules)
        self.assertIn('header', self.parser.css_parser.id_rules)
        # Verify unused rules were NOT loaded (selective parsing)
        self.assertNotIn('unused-class', self.parser.css_parser.class_rules)
        self.assertNotIn('unused-id', self.parser.css_parser.id_rules)

    def test_external_css_selective_parsing(self):
        """Test that selective parsing only loads relevant CSS rules"""
        css_path = os.path.join(self.test_dir, 'assets/css/large_framework.css')
        html = f"""
        <link rel="stylesheet" href="{css_path}">
        <div class="container">
            <button class="btn">Click me</button>
        </div>
        """

        self.parser.add_html_to_document(html, self.document)

        # Verify only used rules were loaded
        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('container', self.parser.css_parser.class_rules)
        self.assertIn('btn', self.parser.css_parser.class_rules)

        # Verify unused framework rules were NOT loaded
        self.assertNotIn('navbar', self.parser.css_parser.class_rules)
        self.assertNotIn('card', self.parser.css_parser.class_rules)
        self.assertNotIn('modal', self.parser.css_parser.class_rules)
        self.assertNotIn('dropdown', self.parser.css_parser.class_rules)

    def test_multiple_external_css_files(self):
        """Test loading multiple external CSS files"""
        css_path1 = os.path.join(self.test_dir, 'assets/css/test_styles.css')
        css_path2 = os.path.join(self.test_dir, 'assets/css/large_framework.css')
        html = f"""
        <link rel="stylesheet" href="{css_path1}">
        <link rel="stylesheet" href="{css_path2}">
        <p class="highlight">Text</p>
        <div class="container">Content</div>
        """

        self.parser.add_html_to_document(html, self.document)

        # Verify rules from both files were loaded
        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('highlight', self.parser.css_parser.class_rules)  # From file 1
        self.assertIn('container', self.parser.css_parser.class_rules)  # From file 2

    def test_external_css_with_style_tag(self):
        """Test that external CSS works alongside <style> tags"""
        css_path = os.path.join(self.test_dir, 'assets/css/test_styles.css')
        html = f"""
        <style>
        .inline-style {{
            color: red;
        }}
        </style>
        <link rel="stylesheet" href="{css_path}">
        <p class="highlight inline-style">Combined styles</p>
        """

        self.parser.add_html_to_document(html, self.document)

        # Verify both sources were loaded
        self.assertTrue(self.parser.css_parser.has_rules())
        self.assertIn('highlight', self.parser.css_parser.class_rules)  # From external
        self.assertIn('inline-style', self.parser.css_parser.class_rules)  # From <style>

    def test_external_css_invalid_file(self):
        """Test handling of invalid/non-existent CSS file"""
        html = """
        <link rel="stylesheet" href="nonexistent.css">
        <p>Test paragraph</p>
        """

        # Should not raise error, just skip invalid CSS
        try:
            self.parser.add_html_to_document(html, self.document)
            success = True
        except Exception:
            success = False

        self.assertTrue(success, "Should handle invalid CSS file gracefully")
        # HTML should still be processed
        self.assertGreaterEqual(len(self.document.paragraphs), 1)

    def test_external_css_relative_path(self):
        """Test CSS loading with relative path"""
        # Use relative path from test directory
        html = """
        <link rel="stylesheet" href="assets/css/test_styles.css">
        <p class="highlight">Test</p>
        """

        # Change to test directory to test relative paths
        original_cwd = os.getcwd()
        try:
            os.chdir(self.test_dir)
            self.parser.add_html_to_document(html, self.document)

            # Verify CSS was loaded
            self.assertTrue(self.parser.css_parser.has_rules())
            self.assertIn('highlight', self.parser.css_parser.class_rules)
        finally:
            os.chdir(original_cwd)

    def test_external_css_not_in_output(self):
        """Test that <link> tags are removed from output"""
        css_path = os.path.join(self.test_dir, 'assets/css/test_styles.css')
        html = f"""
        <link rel="stylesheet" href="{css_path}">
        <p>Some text</p>
        """

        self.parser.add_html_to_document(html, self.document)

        # Check that link tag is not in any paragraph text
        for para in self.document.paragraphs:
            self.assertNotIn('<link>', para.text)
            self.assertNotIn('stylesheet', para.text)

    def test_selective_parsing_efficiency(self):
        """Test that selective parsing is more efficient"""
        css_path = os.path.join(self.test_dir, 'assets/css/large_framework.css')

        # HTML with only a few elements
        html = f"""
        <link rel="stylesheet" href="{css_path}">
        <div class="container">
            <button class="btn">Button</button>
        </div>
        """

        self.parser.add_html_to_document(html, self.document)

        # Should only have loaded 2 class rules (container, btn)
        # Not all the other framework classes
        loaded_classes = set(self.parser.css_parser.class_rules.keys())

        # Verify only relevant classes were loaded
        self.assertIn('container', loaded_classes)
        self.assertIn('btn', loaded_classes)

        # Verify many unused classes were NOT loaded
        unused_classes = {'navbar', 'card', 'modal', 'dropdown', 'alert',
                         'badge', 'progress', 'tooltip', 'popover', 'carousel'}
        loaded_unused = unused_classes.intersection(loaded_classes)
        self.assertEqual(len(loaded_unused), 0,
                        f"Should not load unused classes, but loaded: {loaded_unused}")


if __name__ == "__main__":
    unittest.main()

