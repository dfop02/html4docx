import os
from pathlib import Path
import unittest
from docx import Document
from docx.oxml.ns import qn

from html4docx import HtmlToDocx
from html4docx.utils import unit_converter
from .context import test_dir

class OutputTest(unittest.TestCase):
    @staticmethod
    def clean_up_docx():
        for filename in Path(test_dir).glob("*.docx"):
            filename.unlink()

    @staticmethod
    def get_html_from_file(filename: str):
        file_path = Path(f'{test_dir}/assets/htmls') / Path(filename)
        with open(file_path, 'r') as f:
            html = f.read()
        return html

    @classmethod
    def setUpClass(cls):
        cls.clean_up_docx()
        cls.document = Document()
        cls.text1 = cls.get_html_from_file('text1.html')
        cls.table_html = cls.get_html_from_file('tables1.html')
        cls.table2_html = cls.get_html_from_file('tables2.html')
        cls.table3_html = cls.get_html_from_file('tables3.html')

    @classmethod
    def tearDownClass(cls):
        outputpath = os.path.join(test_dir, 'test.docx')
        cls.document.save(outputpath)

    def setUp(self):
        self.parser = HtmlToDocx()

    def test_html_with_images_links_style(self):
        self.document.add_heading(
            'Test: add regular html with images, links and some formatting to document',
            level=1
        )
        self.parser.add_html_to_document(self.text1, self.document)

    def test_html_with_default_paragraph_style(self):
        self.document.add_heading(
            'Test: add regular html with a default paragraph style defined',
            level=1
        )
        self.parser.paragraph_style = 'Quote'
        self.parser.add_html_to_document(self.text1, self.document)

    def test_add_html_to_table_cell_with_default_paragraph_style(self):
        self.document.add_heading(
            'Test: regular html to table cell with a default paragraph style defined',
            level=1
        )
        self.parser.paragraph_style = 'Quote'
        table = self.document.add_table(2, 2, style='Table Grid')
        cell = table.cell(1, 1)
        self.parser.add_html_to_document(self.text1, cell)

    def test_add_html_to_table_cell(self):
        self.document.add_heading(
            'Test: regular html with images, links, some formatting to table cell',
            level=1
        )
        table = self.document.add_table(2, 2, style='Table Grid')
        cell = table.cell(1, 1)
        self.parser.add_html_to_document(self.text1, cell)

    def test_add_html_skip_images(self):
        self.document.add_heading(
            'Test: regular html with images, but skip adding images',
            level=1
        )
        self.parser.options['images'] = False
        self.parser.add_html_to_document(self.text1, self.document)

        document = self.parser.parse_html_string(self.text1)
        assert any(['Graphic' in paragraph._p.xml for paragraph in document.paragraphs]) is False

    def test_add_html_with_tables(self):
        self.document.add_heading(
            'Test: add html with tables',
            level=1
        )
        self.parser.add_html_to_document(self.table_html, self.document)

        # When no table style is set, use Normal Table as default
        table_style = 'Normal Table'

        # Find the last table added to the document
        last_table = self.document.tables[-1]  # Assumes the table was added at the end

        # Validate the table style
        self.assertEqual(last_table.style.name, table_style, f"Table style does not match expected '{table_style}'")

    def test_add_html_with_tables_accent_style(self):
        table_style = 'Light Grid Accent 6'
        self.document.add_heading(
            'Test: add html with tables with accent',
        )
        self.parser.table_style = table_style
        self.parser.add_html_to_document(self.table_html, self.document)

        # Find the last table added to the document
        last_table = self.document.tables[-1]  # Assumes the table was added at the end

        # Validate the table style
        self.assertEqual(last_table.style.name, table_style, f"Table style does not match expected '{table_style}'")

    def test_add_html_with_tables_basic_style(self):
        table_style = 'Table Grid'
        self.document.add_heading(
            'Test: add html with tables with basic style',
        )
        self.parser.table_style = table_style
        self.parser.add_html_to_document(self.table_html, self.document)

        # Find the last table added to the document
        last_table = self.document.tables[-1]  # Assumes the table was added at the end

        # Validate the table style
        self.assertEqual(last_table.style.name, table_style, f"Table style does not match expected '{table_style}'")

    def test_add_nested_tables(self):
        self.document.add_heading(
            'Test: add nested tables',
        )
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_nested_tables_basic_style(self):
        self.document.add_heading(
            'Test: add nested tables with basic style',
        )
        self.parser.table_style = 'Table Grid'
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_nested_tables_accent_style(self):
        self.document.add_heading(
            'Test: add nested tables with accent style',
        )
        self.parser.table_style = 'Light Grid Accent 6'
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_html_skip_tables(self):
        # broken until feature readded
        self.document.add_heading(
            'Test: add html with tables, but skip adding tables',
            level=1
        )
        self.parser.options['tables'] = False
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_wrong_argument_type_raises_error(self):
        try:
            self.parser.add_html_to_document(self.document, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "First argument needs to be a <class 'str'>" in str(e)
        else:
            assert False, "Error not raised as expected"

        try:
            self.parser.add_html_to_document(self.text1, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "Second argument" in str(e)
            assert "<class 'docx.document.Document'>" in str(e)
        else:
            assert False, "Error not raised as expected"

    def test_add_html_to_cells_method(self):
        self.document.add_heading(
            'Test: add_html_to_cells method',
            level=1
        )
        table = self.document.add_table(2, 3, style='Table Grid')
        cell = table.cell(0, 0)
        html = '''Line 0 without p tags<p>Line 1 with P tags</p>'''
        self.parser.add_html_to_cell(html, cell)

        cell = table.cell(0, 1)
        html = '''<p>Line 0 with p tags</p>Line 1 without p tags'''
        self.parser.add_html_to_cell(html, cell)

        cell = table.cell(0, 2)
        cell.text = "Pre-defined text that shouldn't be removed."
        html = '''<p>Add HTML to non-empty cell.</p>'''
        self.parser.add_html_to_cell(html, cell)

    def test_inline_code(self):
        self.document.add_heading(
            'Test: inline code block',
            level=1
        )

        html = "<p>This is a sentence that contains <code>some code elements</code> that " \
               "should appear as code.</p>"
        self.parser.add_html_to_document(html, self.document)

    def test_code_block(self):
        self.document.add_heading(
            'Test: code block',
            level=1
        )

        html = """<p><code>
This is a code block.
  That should be NOT be pre-formatted.
It should NOT retain carriage returns,

or blank lines.
</code></p>"""
        self.parser.add_html_to_document(html, self.document)

    def test_pre_block(self):
        self.document.add_heading(
            'Test: pre block',
            level=1
        )

        html = """<pre>
This is a pre-formatted block.
  That should be pre-formatted.
Retaining any carriage returns,

and blank lines.
</pre>
"""
        self.parser.add_html_to_document(html, self.document)

    def test_handling_hr(self):
        hr_html_example = '<p>paragraph</p><hr><p>paragraph</p>'

        self.document.add_heading(
            'Test: Handling of hr',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(hr_html_example, self.document)

        document = self.parser.parse_html_string(hr_html_example)
        assert '<w:pBdr>' in document._body._body.xml

    def test_external_hyperlink(self):
        hyperlink_html_example = "<a href=\"https://www.google.com\">Google External Link</a>"

        self.document.add_heading(
            'Test: Handling external hyperlink',
            level=1
        )
        self.parser.add_html_to_document(hyperlink_html_example, self.document)

        document = self.parser.parse_html_string(hyperlink_html_example)
        # Extract external hyperlinks
        external_hyperlinks = []

        for rel in document.part.rels.values():
            if "hyperlink" in rel.reltype:
                external_hyperlinks.append(rel.target_ref)

        assert 'https://www.google.com' in external_hyperlinks
        assert '<w:hyperlink' in document._body._body.xml

    def test_internal_hyperlink(self):
        hyperlink_html_example = (
            "<p><h1 id=\"intro\">Introduction Header</h1></p>"
            "<p>Click here: <a href=\"#intro\" title=\"Link to intro\">Link to intro</a></p>"
        )

        self.document.add_heading(
            'Test: Handling internal hyperlink',
            level=1
        )
        self.parser.add_html_to_document(hyperlink_html_example, self.document)

        document = self.parser.parse_html_string(hyperlink_html_example)
        document_body = document._body._body.xml
        assert '<w:bookmarkStart w:id="0" w:name="intro"/>' in document_body
        assert '<w:bookmarkEnd w:id="0"/>' in document_body
        assert '<w:hyperlink w:anchor="intro" w:tooltip="Link to intro">' in document_body

    def test_image_no_src(self):
        self.document.add_heading(
            'Test: Handling img without src',
            level=1
        )
        self.parser.add_html_to_document('<img />', self.document)

        document = self.parser.parse_html_string('<img />')
        assert '<image: no_src>' in document.paragraphs[0].text

    def test_inline_images(self):
        self.document.add_heading(
            'Test: Handling inline images',
            level=1
        )
        test_img_src = 'https://github.com/dfop02/html4docx/blob/main/tests/assets/images/test_img.png?raw=true'
        html_example = (
            f"<p><img src='{test_img_src}' />"
            f"<img src='{test_img_src}' />"
            f"<img src='{test_img_src}' /></p>"
        )
        self.parser.add_html_to_document(html_example, self.document)

        document = self.parser.parse_html_string(html_example)

        # Find paragraphs containing inline pictures
        img_paragraphs = [
            p for p in document.paragraphs
            if any(r._element.xpath(".//pic:pic") for r in p.runs)
        ]
        assert img_paragraphs, "Expected at least one paragraph with inline images"

        first_img_para = img_paragraphs[0]
        inline_img_runs = [
            r for r in first_img_para.runs
            if r._element.xpath(".//pic:pic")
        ]
        assert len(inline_img_runs) == 3, "Expected 3 inline image runs in a single paragraph"

    def test_bold_italic_underline_and_strike(self):
        self.document.add_heading(
            'Test: Bold, Italic, Underline and Strike tags',
            level=1
        )

        html_example = (
            "<p>This text has <b>Bold Words</b>.</p>"
            "<p>This text has <i>Italic Words</i>.</p>"
            "<p>This text has <u>Underline Words</u>.</p>"
            "<p>This text has <s>Strike Words</s>.</p>"
            "<p>This text has <b><i><u><s>Bold, Italic, Underline and Strike Words</s></u></i></b>.</p>"
        )
        # Add on document for human validation
        self.parser.add_html_to_document(html_example, self.document)

        document = self.parser.parse_html_string(html_example)
        paragraphs = document.paragraphs

        self.assertIn("Bold Words", paragraphs[0].text)
        self.assertTrue(paragraphs[0].runs[2].bold)

        self.assertIn("Italic Words", paragraphs[1].text)
        self.assertTrue(paragraphs[1].runs[2].italic)

        self.assertIn("Underline Words", paragraphs[2].text)
        self.assertTrue(paragraphs[2].runs[2].underline)

        self.assertIn("Strike Words", paragraphs[3].text)
        self.assertTrue(paragraphs[3].runs[2].font.strike)

        self.assertIn("Bold, Italic, Underline and Strike Words", paragraphs[4].text)
        run = paragraphs[4].runs[2]
        self.assertTrue(run.bold)
        self.assertTrue(run.italic)
        self.assertTrue(run.underline)
        self.assertTrue(run.font.strike)

    def test_font_size(self):
        font_size_html_example = (
            "<p><span style=\"font-size:8px\">paragraph 8px</span></p>"
            "<p><span style=\"font-size: 1cm\">paragraph 1cm</span></p>"
            "<p><span style=\"font-size: 6em !important\">paragraph 6em</span></p>"
            "<p><span style=\"font-size: 1.2cm\">paragraph 12cm</span></p>"
            "<p><span style=\"font-size: 1.2vh\">paragraph 12vh not supported</span></p>"
            "<p><span style=\"font-size: 5pc\">paragraph 5pc</span></p>"
            "<p><span style=\"font-size:14pt!important\">paragraph 14pt</span></p>"
            "<p><span style=\"font-size: 16pt!IMPORTANT\">paragraph 16pt</span></p>"
            "<p><span style=\"font-size:2mm!IMPORTANT\">paragraph 2mm</span></p>"
            "<p><span style=\"font-size:small!IMPORTANT\">paragraph small</span></p>"
        )

        self.document.add_heading(
            'Test: Font-Size',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(font_size_html_example, self.document)

        document = self.parser.parse_html_string(font_size_html_example)
        font_sizes = [str(p.runs[1].font.size) for p in document.paragraphs]
        assert ['76200', '355600', '914400', '431800', 'None', '762000', '177800', '203200', '69850', '120650'] == font_sizes

    def test_color_by_name(self):
        color_html_example = (
            "<p><span style=\"color:red\">paragraph red</span></p>"
            "<p><span style=\"color: yellow\">paragraph yellow</span></p>"
            "<p><span style=\"color: blue !important\">paragraph blue</span></p>"
            "<p><span style=\"color: green!important\">paragraph green</span></p>"
            "<p><span style=\"color: darkgray!IMPORTANT\">paragraph darkgray</span></p>"
            "<p><span style=\"color: MAGENTA !IMPORTANT\">paragraph magenta</span></p>"
            "<p><span style=\"color: invalidcolor\">paragraph has default black because of invalid color name</span></p>"
        )

        self.document.add_heading(
            'Test: Color by name',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(color_html_example, self.document)

        document = self.parser.parse_html_string(color_html_example)
        colors = [str(p.runs[1].font.color.rgb) for p in document.paragraphs]

        assert 'FF0000' in colors # Red
        assert 'FFFF00' in colors # Yellow
        assert '0000FF' in colors # Blue
        assert '008000' in colors # Green
        assert 'A9A9A9' in colors # Darkgray
        assert '000000' in colors # Black
        assert 'FF00FF' in colors # Magenta

    def test_table_cell_border_properties(self):
        """Validates that all table cells have the expected border size, style, and color."""

        self.document.add_heading(
            'Test: Table Cell Border Properties',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(self.table3_html, self.document)
        document = self.parser.parse_html_string(self.table3_html)

        # Define expected border properties
        expected_colors = [
            {
                "top": {"color": "D95B48", "style": "single", "size": "1.0"},
                "bottom": {"color": "D95B48", "style": "single", "size": "1.0"},
                "left": {"color": "FF0000", "style": "single", "size": "1.0"},
                "right": {"color": "8B0000", "style": "single", "size": "1.0"}
            },
            {
                "top": {"color": "FAC32A", "style": "single", "size": "1.0"},
                "bottom": {"color": "FAC32A", "style": "single", "size": "1.0"},
                "left": {"color": "none", "style": "none", "size": "none"},
                "right": {"color": "FAC32A", "style": "single", "size": "12.0"}
            },
            {
                "top": {"color": "30E667", "style": "none", "size": "5.0"},
                "bottom": {"color": "30E667", "style": "single", "size": "5.0"},
                "left": {"color": "30E667", "style": "single", "size": "5.0"},
                "right": {"color": "30E667", "style": "single", "size": "5.0"}
            },
            {
                "top": {"color": "none", "style": "none", "size": "none"},
                "bottom": {"color": "D948CF", "style": "single", "size": "1.0"},
                "left": {"color": "none", "style": "none", "size": "none"},
                "right": {"color": "D948CF", "style": "single", "size": "5.0"}
            },
            {
                "top": {"color": "EAAAA7", "style": "single", "size": "1.0"},
                "bottom": {"color": "EAAAA7", "style": "single", "size": "1.0"},
                "left": {"color": "EAAAA7", "style": "single", "size": "1.0"},
                "right": {"color": "EAAAA7", "style": "single", "size": "1.0"}
            },
            {
                "top": {"color": "none", "style": "none", "size": "none"},
                "bottom": {"color": "ACC4AA", "style": "dashed", "size": "7.0"},
                "left": {"color": "none", "style": "none", "size": "none"},
                "right": {"color": "ACC4AA", "style": "dotted", "size": "4.0"}
            }
        ]

        # Validate border properties for each cell
        cell_idx = 0
        for row_idx, row in enumerate(document.tables[0].rows):
            for column_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))

                # Extract border properties
                border_sides = {
                    'top': tcBorders.find(qn('w:top')) if tcBorders is not None else None,
                    'bottom': tcBorders.find(qn('w:bottom')) if tcBorders is not None else None,
                    'left': tcBorders.find(qn('w:left')) if tcBorders is not None else None,
                    'right': tcBorders.find(qn('w:right')) if tcBorders is not None else None,
                }

                for side, border in border_sides.items():
                    if border is not None:
                        color = border.get(qn('w:color'), "").upper()  # Ensure uppercase and no #
                        size = border.get(qn('w:sz'))
                        style = border.get(qn('w:val'))
                    else:
                        color, size, style = "none", "none", "none"

                    # Convert size from eighths of a point to points
                    size_in_pt = str(int(size) / 8) if size and size != "none" else "none"

                    # Get expected properties for the current cell and side
                    expected_properties = expected_colors[cell_idx][side]

                    # Assertions
                    assert color == expected_properties["color"], (
                        f"Color mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['color']}, got {color}"
                    )
                    assert size_in_pt == expected_properties["size"], (
                        f"Size mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['size']}, got {size_in_pt}"
                    )
                    assert style == expected_properties["style"], (
                        f"Style mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['style']}, got {style}"
                    )

                cell_idx += 1

    def test_table_cell_background_color(self):
        """Validates that all table cells have the expected background color."""

        self.document.add_heading(
            'Test: Table Cell Background Color',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(self.table3_html, self.document)
        document = self.parser.parse_html_string(self.table3_html)

        # Define expected background colors for each cell
        expected_background_colors = [
            "3749EF", # Row 1 Column 1
            "33b32e", # Row 1 Column 2
            "BFBFBF", # Row 2 Column 1
            "2eaab3", # Row 2 Column 2
            "99fffa", # Row 3 Column 1
            "2eaab3"  # Row 3 Column 2
        ]

        # Validate background colors for each cell
        cell_idx = 0
        for row_idx, row in enumerate(document.tables[0].rows):
            for column_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Get the background color (shading) if it exists
                shading = tcPr.find(qn('w:shd'))
                if shading is not None:
                    background_color = shading.get(qn('w:fill'), "").upper()  # Ensure uppercase and no #
                else:
                    background_color = "None"

                # Get expected background color for the current cell
                expected_color = expected_background_colors[cell_idx].upper()
                cell_idx += 1

                assert background_color == expected_color, (
                    f"Background color mismatch for row {row_idx} column {column_idx}: "
                    f"expected {expected_color}, got {background_color}"
                )

    def test_table_cell_dimensions(self):
        """Validates that all table cells have the expected width and height."""

        self.document.add_heading(
            'Test: Table Cell Dimensions',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(self.table3_html, self.document)
        document = self.parser.parse_html_string(self.table3_html)

        # Define expected dimensions for each cell
        expected_dimensions = [
            # First row
            [
                {
                    "width": "258.35px",  # Width for the first cell
                    "height": "23.75pt"   # Height for the first cell
                },
                {
                    "width": "222.2pt",   # Width for the second cell
                    "height": "23.75pt"   # Height for the second cell
                }
            ],
            # Second row
            [
                {
                    "width": "258.35in",  # Width for the first cell
                    "height": "15.5pt"    # Height for the first cell
                },
                {
                    "width": "6cm",       # Width for the second cell
                    "height": "15.5pt"    # Height for the second cell
                }
            ],
            # Third row
            [
                {
                    "width": "258.35pt",  # Width for the first cell
                    "height": "2rem"      # Height for the first cell
                },
                {
                    "width": "6cm",       # Width for the second cell
                    "height": "2em"       # Height for the second cell
                }
            ]
        ]

        # Validate dimensions for each cell
        for row_idx, row in enumerate(document.tables[0].rows):
            for cell_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                docx_cell = document.tables[0].cell(row_idx, cell_idx)

                # Convert width from EMUs to px
                cell_width_px = round((docx_cell.width / 914400) * 96, 2)  # 1 EMU = 1/914400 inch, 1 inch = 96px
                # Get expected width and convert it to points using unit_converter
                expected_width = expected_dimensions[row_idx][cell_idx]["width"]
                expected_width_px = unit_converter(expected_width, "px")

                # Convert height from EMUs to px
                cell_height_px = round((row.height / 914400) * 96, 2)  # 1 EMU = 1/914400 inch, 1 inch = 96px
                # Get expected height and convert it to points
                expected_height = expected_dimensions[row_idx][cell_idx]["height"]
                expected_height_px = unit_converter(expected_height, "px")

                assert round(abs(cell_width_px - expected_width_px), 2) <= 0.03, (
                    f"Width mismatch for cell ({row_idx}, {cell_idx}): "
                    f"expected {expected_width_px}px, got {cell_width_px}px"
                )
                assert round(abs(cell_height_px - expected_height_px), 2) <= 0.03, (
                    f"Height mismatch for cell ({row_idx}, {cell_idx}): "
                    f"expected {expected_height_px}px, got {cell_height_px}px"
                )

    def test_unbalanced_table(self):
        # A table with more td elements in latter rows than in the first
        self.document.add_heading(
            'Test: Handling unbalanced tables',
            level=1
        )
        self.parser.add_html_to_document(
            "<table>"
            "<tr><td>Hello</td></tr>"
            "<tr><td>One</td><td>Two</td></tr>"
            "</table>",
            self.document
        )

    def test_emojis_and_special_characters(self):
        emojis_and_special_chars_html_example = """
        <html>
            <body>
                <p>Emoji Test: 😊🔥🎉</p>
                <p>HTML Entities: &amp; &lt; &gt; &copy; &reg;</p>
                <p>Math Symbols: ∑ π √</p>
            </body>
        </html>
        """

        self.document.add_heading(
            'Test: Emojis and Special Characters',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(emojis_and_special_chars_html_example, self.document)

        document = self.parser.parse_html_string(emojis_and_special_chars_html_example)
        doc_text = " ".join([p.text for p in document.paragraphs])

        # Check if all expected elements exist in the DOCX
        assert "😊" in doc_text, "Emoji '😊' is missing"
        assert "🔥" in doc_text, "Emoji '🔥' is missing"
        assert "🎉" in doc_text, "Emoji '🎉' is missing"
        assert "&" in doc_text, "HTML entity '&amp;' did not convert correctly"
        assert "<" in doc_text, "HTML entity '&lt;' did not convert correctly"
        assert ">" in doc_text, "HTML entity '&gt;' did not convert correctly"
        assert "©" in doc_text, "HTML entity '&copy;' did not convert correctly"
        assert "®" in doc_text, "HTML entity '&reg;' did not convert correctly"
        assert "∑" in doc_text, "Math symbol '∑' is missing"
        assert "π" in doc_text, "Math symbol 'π' is missing"
        assert "√" in doc_text, "Math symbol '√' is missing"

    def test_ordered_list(self):
        self.document.add_heading("Test: Ordered List", level=1)

        ordered_list_html_example = """
        <ol>
          <li>first list, item 1</li>
          <p>Paragraph inserted between items</p>
          <li>first list, item 2</li>
          <li><p>first list, item 3 within a paragraph</p></li>
        </ol>
        """
        self.parser.add_html_to_document(ordered_list_html_example, self.document)
        document = self.parser.parse_html_string(ordered_list_html_example)

        # Extract paragraphs with 'ListNumber' style (ordered list)
        ordered_list_paragraphs = [
            p for p in document.paragraphs if p.style.name == "List Number"
        ]

        # Expected items in order
        expected_items = [
            "first list, item 1",
            "first list, item 2",
            "first list, item 3 within a paragraph",
        ]

        assert len(ordered_list_paragraphs) >= len(
            expected_items
        ), f"Expected at least {len(expected_items)} ordered list items, found {len(ordered_list_paragraphs)}"

        for i, expected_text in enumerate(expected_items):
            actual_text = ordered_list_paragraphs[i].text.strip()
            assert (
                actual_text == expected_text
            ), f"Expected ordered list item '{expected_text}', but got '{actual_text}'"

    def test_unordered_list(self):
        self.document.add_heading("Test: Unordered List", level=1)

        unordered_list_html_example = """
        <ul>
          <li>Unorderd list</li>
          <p>Paragraph inserted between items</p>
          <li>with circle markers</li>
          <li><p>last option</p></li>
        </ul>
        """
        self.parser.add_html_to_document(unordered_list_html_example, self.document)
        document = self.parser.parse_html_string(unordered_list_html_example)

        # Extract paragraphs with 'ListBullet' style (unordered list)
        unordered_list_paragraphs = [
            p for p in document.paragraphs if p.style.name == "List Bullet"
        ]

        # Expected unordered items
        expected_items = ["Unorderd list", "with circle markers", "last option"]

        assert len(unordered_list_paragraphs) >= len(
            expected_items
        ), f"Expected at least {len(expected_items)} unordered list items, found {len(unordered_list_paragraphs)}"

        for expected_text in expected_items:
            assert any(
                expected_text == p.text.strip() for p in unordered_list_paragraphs
            ), f"Unordered list item '{expected_text}' not found in List Bullet paragraphs"


if __name__ == "__main__":
    unittest.main()
