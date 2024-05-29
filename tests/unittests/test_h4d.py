"""Unit test suite that covers html4docx/h4d.py."""

import os
import pathlib
from html4docx.h4d import HtmlToDocx
from docx import Document


def test_h4d_set_table_style() -> None:
    """Test that covers set_table_style() in h4d.py."""
    test_docx_filename = "test_h4d_set_table_style.docx"

    with open(f"{str(pathlib.Path(__file__).parent.resolve())}/heyo.html", "r", encoding="utf-8") as heyo_html:
        test_html_string = heyo_html.read()

    # Test where table_style and table.style are both None.
    parser = HtmlToDocx()
    document = Document()
    parser.add_html_to_document(test_html_string, document)
    document.save(test_docx_filename)
    assert not parser.table_style
    assert hasattr(parser, "table")
    assert not parser.table

    # Test where table_style is set to valid table Style without spaces but parsed to correct syntax.
    parser = HtmlToDocx()
    parser.set_table_style("TableGrid")
    assert parser.table_style == "Table Grid"
    document = Document()
    parser.add_html_to_document(test_html_string, document)
    document.save(test_docx_filename)
    assert parser.table_style == "Table Grid"

    # Test where table_style is set to valid table Style with spaces.
    parser = HtmlToDocx()
    parser.set_table_style("Colorful Grid Accent 2")
    assert parser.table_style == "Colorful Grid Accent 2"
    document = Document()
    parser.add_html_to_document(test_html_string, document)
    document.save(test_docx_filename)
    assert parser.table_style == "Colorful Grid Accent 2"

    # Cleanup test file.
    os.remove(test_docx_filename)
