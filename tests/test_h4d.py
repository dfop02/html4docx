import os
import unittest
from io import BytesIO
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from html4docx import HtmlToDocx
from html4docx.utils import unit_converter, parse_color
from html4docx.colors import Color

test_dir = os.path.abspath(os.path.dirname(__file__))

class OutputTest(unittest.TestCase):
    # ============================== Helper methods ============================== #
    @staticmethod
    def clean_up_docx():
        for filename in Path(test_dir).glob("*.docx"):
            filename.unlink()

    @staticmethod
    def get_html_from_file(filename: str):
        file_path = Path(f'{test_dir}/assets/htmls') / Path(filename)
        with open(file_path, 'r') as f:
            html = f.read()
        return html

    @staticmethod
    def hexcolor(color: str) -> str:
        """
        Convert a color string to a hex string.
        Returns a hex string like 'FF0000'.
        """
        return parse_color(color, return_hex=True).lstrip("#")

    @staticmethod
    def get_underline_color(run):
        """
        Extract underline color from the run XML.
        Returns hex string like 'FF0000' or None.
        """
        u_elems = run._r.xpath('.//w:u')
        if not u_elems:
            return None

        u_elem = u_elems[0]
        return u_elem.get(qn('w:color'))

    # ============================== Setup and teardown ============================== #
    @classmethod
    def setUpClass(cls):
        cls.clean_up_docx()
        cls.document = Document()
        cls.text1 = cls.get_html_from_file('text1.html')
        cls.paragraph_line_height = cls.get_html_from_file('paragraph_line_height.html')
        cls.paragraph_first_line_indent = cls.get_html_from_file('paragraph_first_line_indent.html')
        cls.text_decoration = cls.get_html_from_file('text_decoration.html')
        cls.css_properties = cls.get_html_from_file('css_properties.html')
        cls.css_properties_header = cls.get_html_from_file('header.html')
        cls.table_html = cls.get_html_from_file('tables1.html')
        cls.table2_html = cls.get_html_from_file('tables2.html')
        cls.table3_html = cls.get_html_from_file('tables3.html')

    @classmethod
    def tearDownClass(cls):
        outputpath = os.path.join(test_dir, 'test.docx')
        cls.document.save(outputpath)

    def setUp(self):
        self.parser = HtmlToDocx()

    # ============================== Tests ============================== #
    def test_save_docx_by_filename(self):
        filename = os.path.join(test_dir, 'new_test.docx')
        self.parser.set_initial_attrs(self.document)
        self.parser.save(filename)
        self.assertTrue(os.path.exists(filename))
        os.remove(filename)

    def test_save_docx_by_buffer(self):
        buffer = BytesIO()
        self.parser.set_initial_attrs(self.document)
        self.parser.save(buffer)
        buffer.seek(0)
        self.assertTrue(buffer.getvalue())

    def test_html_with_images_links_style(self):
        self.document.add_heading(
            'Test: add regular html with images, links and some formatting to document',
            level=1
        )
        self.parser.add_html_to_document(self.text1, self.document)

    def test_html_with_default_paragraph_style(self):
        self.document.add_heading(
            'Test: add regular html with a default paragraph style defined',
            level=1
        )
        self.parser.paragraph_style = 'Quote'
        self.parser.add_html_to_document(self.text1, self.document)

    def test_add_html_to_table_cell_with_default_paragraph_style(self):
        self.document.add_heading(
            'Test: regular html to table cell with a default paragraph style defined',
            level=1
        )
        self.parser.paragraph_style = 'Quote'
        table = self.document.add_table(1, 2, style='Table Grid')
        cell = table.cell(0, 1)
        self.parser.add_html_to_document(self.text1, cell)

    def test_add_html_to_table_cell(self):
        self.document.add_heading(
            'Test: regular html with images, links, some formatting to table cell',
            level=1
        )
        table = self.document.add_table(1, 2, style='Table Grid')
        cell = table.cell(0, 1)
        self.parser.add_html_to_document(self.text1, cell)

    def test_add_html_skip_images(self):
        self.document.add_heading(
            'Test: regular html with images, but skip adding images',
            level=1
        )
        self.parser.options['images'] = False
        self.parser.add_html_to_document(self.text1, self.document)

        document = self.parser.parse_html_string(self.text1)
        assert any(['Graphic' in paragraph._p.xml for paragraph in document.paragraphs]) is False

    def test_add_html_with_tables(self):
        self.document.add_heading(
            'Test: add html with tables (by default no borders)',
            level=1
        )
        self.parser.add_html_to_document(self.table_html, self.document)

        # When no table style is set, use Normal Table as default
        table_style = 'Normal Table'

        # Find the last table added to the document
        last_table = self.document.tables[-1]  # Assumes the table was added at the end

        # Validate the table style
        self.assertEqual(last_table.style.name, table_style, f"Table style does not match expected '{table_style}'")

    def test_add_html_with_tables_accent_style(self):
        table_style = 'Light Grid Accent 6'
        self.document.add_heading(
            'Test: add html with tables with accent',
        )
        self.parser.table_style = table_style
        self.parser.add_html_to_document(self.table_html, self.document)

        # Find the last table added to the document
        last_table = self.document.tables[-1]  # Assumes the table was added at the end

        # Validate the table style
        self.assertEqual(last_table.style.name, table_style, f"Table style does not match expected '{table_style}'")

    def test_add_html_with_tables_basic_style(self):
        table_style = 'Table Grid'
        self.document.add_heading(
            'Test: add html with tables with basic style',
        )
        self.parser.table_style = table_style
        self.parser.add_html_to_document(self.table_html, self.document)

        # Find the last table added to the document
        last_table = self.document.tables[-1]  # Assumes the table was added at the end

        # Validate the table style
        self.assertEqual(last_table.style.name, table_style, f"Table style does not match expected '{table_style}'")

    def test_add_nested_tables(self):
        self.document.add_heading(
            'Test: add nested tables',
        )
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_nested_tables_basic_style(self):
        self.document.add_heading(
            'Test: add nested tables with basic style',
        )
        self.parser.table_style = 'Table Grid'
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_nested_tables_accent_style(self):
        self.document.add_heading(
            'Test: add nested tables with accent style',
        )
        self.parser.table_style = 'Light Grid Accent 6'
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_html_skip_tables(self):
        # broken until feature readded
        self.document.add_heading(
            'Test: add html with tables, but skip adding tables',
            level=1
        )
        self.parser.options['tables'] = False
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_wrong_argument_type_raises_error(self):
        try:
            self.parser.add_html_to_document(self.document, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "First argument needs to be a <class 'str'>" in str(e)
        else:
            assert False, "Error not raised as expected"

        try:
            self.parser.add_html_to_document(self.text1, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "Second argument" in str(e)
            assert "<class 'docx.document.Document'>" in str(e)
        else:
            assert False, "Error not raised as expected"

    def test_add_html_to_cells_method(self):
        self.document.add_heading(
            'Test: add_html_to_cells method',
            level=1
        )
        table = self.document.add_table(2, 3, style='Table Grid')
        cell = table.cell(0, 0)
        html = '''Line 0 without p tags<p>Line 1 with P tags</p>'''
        self.parser.add_html_to_cell(html, cell)

        cell = table.cell(0, 1)
        html = '''<p>Line 0 with p tags</p>Line 1 without p tags'''
        self.parser.add_html_to_cell(html, cell)

        cell = table.cell(0, 2)
        cell.text = "Pre-defined text that shouldn't be removed."
        html = '''<p>Add HTML to non-empty cell.</p>'''
        self.parser.add_html_to_cell(html, cell)

    def test_inline_code(self):
        self.document.add_heading(
            'Test: inline code block',
            level=1
        )

        html = "<p>This is a sentence that contains <code>some code elements</code> that " \
               "should appear as code.</p>"
        self.parser.add_html_to_document(html, self.document)

    def test_code_block(self):
        self.document.add_heading(
            'Test: code block',
            level=1
        )

        html = """<p><code>
This is a code block.
  That should be NOT be pre-formatted.
It should NOT retain carriage returns,

or blank lines.
</code></p>"""
        self.parser.add_html_to_document(html, self.document)

    def test_pre_block(self):
        self.document.add_heading(
            'Test: pre block',
            level=1
        )

        html = """<pre>
This is a pre-formatted block.
  That should be pre-formatted.
Retaining any carriage returns,

and blank lines.
</pre>
"""
        self.parser.add_html_to_document(html, self.document)

    def test_handling_hr(self):
        hr_html_example = '<p>paragraph</p><hr><p>paragraph</p>'

        self.document.add_heading(
            'Test: Handling of hr',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(hr_html_example, self.document)

        document = self.parser.parse_html_string(hr_html_example)
        assert '<w:pBdr>' in document._body._body.xml

    def test_external_hyperlink(self):
        hyperlink_html_example = "<a href=\"https://www.google.com\">Google External Link</a>"

        self.document.add_heading(
            'Test: Handling external hyperlink',
            level=1
        )
        self.parser.add_html_to_document(hyperlink_html_example, self.document)

        document = self.parser.parse_html_string(hyperlink_html_example)
        # Extract external hyperlinks
        external_hyperlinks = []

        for rel in document.part.rels.values():
            if "hyperlink" in rel.reltype:
                external_hyperlinks.append(rel.target_ref)

        assert 'https://www.google.com' in external_hyperlinks
        assert '<w:hyperlink' in document._body._body.xml

    def test_internal_hyperlink(self):
        hyperlink_html_example = (
            "<p><h1 id=\"intro\">Introduction Header</h1></p>"
            "<p>Click here: <a href=\"#intro\" title=\"Link to intro\">Link to intro</a></p>"
        )

        self.document.add_heading(
            'Test: Handling internal hyperlink',
            level=1
        )
        self.parser.add_html_to_document(hyperlink_html_example, self.document)

        document = self.parser.parse_html_string(hyperlink_html_example)
        document_body = document._body._body.xml
        assert '<w:bookmarkStart w:id="0" w:name="intro"/>' in document_body
        assert '<w:bookmarkEnd w:id="0"/>' in document_body
        assert '<w:hyperlink w:anchor="intro" w:tooltip="Link to intro">' in document_body

    def test_internal_hyperlink_without_paragraph(self):
        hyperlink_html_example = (
            "<h1 id=\"intro\">Introduction Header</h1>"
            "<p>Click here: <a href=\"#intro\" title=\"Link to intro\">Link to intro</a><p/>"
        )

        document = self.parser.parse_html_string(hyperlink_html_example)
        document_body = document._body._body.xml

        assert '<w:bookmarkStart w:id="0" w:name="intro"/>' in document_body
        assert '<w:bookmarkEnd w:id="0"/>' in document_body
        assert '<w:hyperlink w:anchor="intro" w:tooltip="Link to intro">' in document_body

    def test_internal_hyperlink_without_anchor(self):
        hyperlink_html_example = (
            "<p>Click here: <a href=\"#intro\" title=\"Link to intro\">Link to intro</a></p>"
        )

        document = self.parser.parse_html_string(hyperlink_html_example)
        document_body = document._body._body.xml

        assert '<w:bookmarkStart w:id="0" w:name="intro"/>' not in document_body
        assert '<w:bookmarkEnd w:id="0"/>' not in document_body
        assert '<w:hyperlink w:anchor="intro" w:tooltip="Link to intro">' in document_body

    def test_image_no_src(self):
        self.document.add_heading(
            'Test: Handling img without src',
            level=1
        )
        self.parser.add_html_to_document('<img />', self.document)

        document = self.parser.parse_html_string('<img />')
        assert '<image: no_src>' in document.paragraphs[0].text

    def test_local_img(self):
        # A table with more td elements in latter rows than in the first
        self.document.add_heading('Test: Local Image', level=1)
        html_local_img = '<img alt="" height="306px" src="./tests/assets/images/test_img.png" width="520px"/>'
        self.parser.add_html_to_document(html_local_img, self.document)
        document = self.parser.parse_html_string(html_local_img)

        # Get the last paragraph
        paragraphs = document.paragraphs
        image_paragraph = paragraphs[-1]

        # Check the run contains an image
        image_found = False
        for run in image_paragraph.runs:
            if run._element.xpath(".//w:drawing"):
                image_found = True
                break

        assert image_found, "No image was found in the document"

    def test_inline_images(self):
        self.document.add_heading(
            'Test: Handling inline images',
            level=1
        )
        test_img_src = 'https://github.com/dfop02/html4docx/blob/main/tests/assets/images/test_img.png?raw=true'
        html_example = (
            f"<p><img src='{test_img_src}' />"
            f"<img src='{test_img_src}' />"
            f"<img src='{test_img_src}' /></p>"
        )
        self.parser.add_html_to_document(html_example, self.document)

        document = self.parser.parse_html_string(html_example)

        # Find paragraphs containing inline pictures
        img_paragraphs = [
            p for p in document.paragraphs
            if any(r._element.xpath(".//pic:pic") for r in p.runs)
        ]
        assert img_paragraphs, "Expected at least one paragraph with inline images"

        first_img_para = img_paragraphs[0]
        inline_img_runs = [
            r for r in first_img_para.runs
            if r._element.xpath(".//pic:pic")
        ]
        assert len(inline_img_runs) == 3, "Expected 3 inline image runs in a single paragraph"

    def test_single_image_without_paragraph(self):
        html_example = "<img src='https://github.com/dfop02/html4docx/blob/main/tests/assets/images/test_img.png?raw=true' />"
        document = self.parser.parse_html_string(html_example)

        # Find paragraphs containing inline pictures
        img_paragraphs = [
            p for p in document.paragraphs
            if any(r._element.xpath(".//pic:pic") for r in p.runs)
        ]
        assert img_paragraphs, "Expected at least one paragraph with inline images"

        first_img_para = img_paragraphs[0]
        inline_img_runs = [
            r for r in first_img_para.runs
            if r._element.xpath(".//pic:pic")
        ]
        assert len(inline_img_runs) == 1, "Expected 1 inline image runs in a single paragraph"

    def test_bold_italic_underline_and_strike(self):
        self.document.add_heading(
            'Test: Bold, Italic, Underline and Strike tags',
            level=1
        )

        html_example = (
            "<p>This text has <b>Bold Words</b>.</p>"
            "<p>This text has <i>Italic Words</i>.</p>"
            "<p>This text has <u>Underline Words</u>.</p>"
            "<p>This text has <s>Strike Words</s>.</p>"
            "<p>This text has <b><i><u><s>Bold, Italic, Underline and Strike Words</s></u></i></b>.</p>"
        )
        # Add on document for human validation
        self.parser.add_html_to_document(html_example, self.document)

        document = self.parser.parse_html_string(html_example)
        paragraphs = document.paragraphs

        self.assertIn("Bold Words", paragraphs[0].text)
        self.assertTrue(paragraphs[0].runs[1].bold)

        self.assertIn("Italic Words", paragraphs[1].text)
        self.assertTrue(paragraphs[1].runs[1].italic)

        self.assertIn("Underline Words", paragraphs[2].text)
        self.assertTrue(paragraphs[2].runs[1].underline)

        self.assertIn("Strike Words", paragraphs[3].text)
        self.assertTrue(paragraphs[3].runs[1].font.strike)

        self.assertIn("Bold, Italic, Underline and Strike Words", paragraphs[4].text)
        run = paragraphs[4].runs[1]
        self.assertTrue(run.bold)
        self.assertTrue(run.italic)
        self.assertTrue(run.underline)
        self.assertTrue(run.font.strike)

    def test_font_size(self):
        font_size_html_example = (
            "<p><span style=\"font-size:8px\">paragraph 8px</span></p>"
            "<p><span style=\"font-size: 1cm\">paragraph 1cm</span></p>"
            "<p><span style=\"font-size: 6em !important\">paragraph 6em</span></p>"
            "<p><span style=\"font-size: 1.2cm\">paragraph 12cm</span></p>"
            "<p><span style=\"font-size: 1.2vh\">paragraph 12vh not supported</span></p>"
            "<p><span style=\"font-size: 5pc\">paragraph 5pc</span></p>"
            "<p><span style=\"font-size:14pt!important\">paragraph 14pt</span></p>"
            "<p><span style=\"font-size: 16pt!IMPORTANT\">paragraph 16pt</span></p>"
            "<p><span style=\"font-size:2mm!IMPORTANT\">paragraph 2mm</span></p>"
            "<p><span style=\"font-size:small!IMPORTANT\">paragraph small</span></p>"
        )

        self.document.add_heading(
            'Test: Font-Size',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(font_size_html_example, self.document)

        document = self.parser.parse_html_string(font_size_html_example)
        font_sizes = [str(p.runs[0].font.size) for p in document.paragraphs]
        assert ['76200', '355600', '914400', '431800', 'None', '762000', '177800', '203200', '69850', '120650'] == font_sizes

    def test_font_size_paragraph(self):
        font_size_html_example = (
            "<p style=\"font-size:8px\">paragraph 8px</p>"
            "<p style=\"font-size: 1cm\">paragraph 1cm</p>"
            "<p style=\"font-size: 6em !important\">paragraph 6em</p>"
            "<p style=\"font-size: 1.2cm\">paragraph 12cm</p>"
            "<p style=\"font-size: 1.2vh\">paragraph 12vh not supported</p>"
            "<p style=\"font-size: 5pc\">paragraph 5pc</p>"
            "<p style=\"font-size:14pt!important\">paragraph 14pt</p>"
            "<p style=\"font-size: 16pt!IMPORTANT\">paragraph 16pt</p>"
            "<p style=\"font-size:2mm!IMPORTANT\">paragraph 2mm</p>"
            "<p style=\"font-size:small!IMPORTANT\">paragraph small</p>"
        )

        self.document.add_heading(
            'Test: Font-Size on <p>',
            level=1
        )
        self.parser.add_html_to_document(font_size_html_example, self.document)

        document = self.parser.parse_html_string(font_size_html_example)
        font_sizes = [str(p.runs[0].font.size) for p in document.paragraphs]
        assert ['76200', '355600', '914400', '431800', 'None', '762000', '177800', '203200', '69850', '120650'] == font_sizes

    def test_font_weight_paragraph(self):
        self.document.add_heading('Test: font weight on <p>', level=1)
        font_weight_html_example = (
            "<p style=\"font-weight: bold\">bold text</p>"
            "<p style=\"font-weight: bolder\">bolder text</p>"
            "<p style=\"font-weight: 700\">700 weight</p>"
            "<p style=\"font-weight: 900\">900 weight</p>"
            "<p style=\"font-weight: normal\">normal text</p>"
            "<p style=\"font-weight: lighter\">lighter text</p>"
            "<p style=\"font-weight: 400\">400 weight</p>"
            "<p style=\"font-weight: 100\">100 weight</p>"
        )

        self.parser.add_html_to_document(font_weight_html_example, self.document)

        document = self.parser.parse_html_string(font_weight_html_example)

        font_weights = [p.runs[0].font.bold for p in document.paragraphs]

        expected_weights = [
            True,   # bold
            True,   # bolder
            True,   # 700
            True,   # 900
            False,  # normal
            False,  # lighter
            False,  # 400
            False,  # 100
        ]

        self.assertEqual(font_weights, expected_weights)

    def test_font_style_paragraph(self):
        self.document.add_heading('Test: font style on <p>', level=1)
        font_style_html_example = (
            "<p style=\"font-style: italic\">italic text</p>"
            "<p style=\"font-style: oblique\">oblique text</p>"
            "<p style=\"font-style: normal\">normal text</p>"
        )

        self.parser.add_html_to_document(font_style_html_example, self.document)

        document = self.parser.parse_html_string(font_style_html_example)

        font_styles = [p.runs[0].font.italic for p in document.paragraphs]

        expected_styles = [
            True,   # italic
            True,   # oblique (should be treated as italic)
            False,  # normal
        ]

        self.assertEqual(font_styles, expected_styles)

    def test_font_family_paragraph(self):
        self.document.add_heading('Test: font family on <p>', level=1)
        font_family_html_example = (
            "<p style=\"font-family: Arial, sans-serif\">Arial font text</p>"
            "<p style=\"font-family: 'Helvetica', sans-serif\">Helvetica font text</p>"
            "<p style=\"font-family: 'Noto Sans', sans-serif\">Noto Sans font text</p>"
            "<p style=\"font-family: 'Times New Roman', serif\">Times New Roman font text</p>"
            "<p style=\"font-family: serif\">Generic serif font text</p>"
            "<p style=\"font-family: sans-serif\">Generic sans-serif font text</p>"
            "<p style=\"font-family: monospace\">Generic monospace font text</p>"
            "<p style=\"font-family: 'Courier New', monospace\">Courier New font text</p>"
            "<p style=\"font-family: inherit\">Inherit font text</p>"
        )

        self.parser.add_html_to_document(font_family_html_example, self.document)

    def test_text_transform_paragraph(self):
        self.document.add_heading('Test: text-transform on <p>', level=1)
        text_transform_html_example = (
            "<p style=\"text-transform: uppercase\">uppercase text</p>"
            "<p style=\"text-transform: lowercase\">LOWERCASE TEXT</p>"
            "<p style=\"text-transform: capitalize\">capitalize each word</p>"
            "<p style=\"text-transform: none\">normal text</p>"
            "<p>default text</p>"
        )

        self.parser.add_html_to_document(text_transform_html_example, self.document)

    def test_text_decoration_span(self):
        self.document.add_heading('Test: text-decoration on <span>', level=1)
        text_decoration_html_example = (
            # Standalone spans
            "<span style=\"text-decoration: underline red\">underlined span (red)</span>"
            "<span style=\"text-decoration: none rgb(0,0,0)\">no decoration span (rgb(0, 0, 0))</span>"
            "<span style=\"text-decoration: line-through gray\">strikethrough span (gray) (not supported)</span>"
            "<span style=\"text-decoration: underline line-through orange\">underline+line-through span (orange)\
                (should be strike)</span>"

            # Spans inside paragraphs
            "<p>Normal text <span style=\"text-decoration: underline wavy blue\">wavy underlined span (blue)</span> continues</p>"
            "<p>Normal text <span style=\"text-decoration: underline dotted purple\">dotted underlined span (purple)</span> continues</p>"
            "<p>Normal text <span style=\"text-decoration: line-through red\">strikethrough span (red)</span> continues</p>"

            # Multiple spans with different decorations in same paragraph
            "<p>Start <span style=\"text-decoration: underline green\">underlined</span> "
            "<span style=\"text-decoration: line-through blue\">strikethrough</span> "
            "<span style=\"text-decoration: underline dashed orange\">dashed underline</span> end</p>"

            # Span with no decoration inside decorated paragraph
            "<p style=\"text-decoration: underline\">Underlined paragraph with "
            "<span style=\"text-decoration: none\">normal span</span> inside</p>"

            # Span with decoration inside decorated paragraph (should override)
            "<p style=\"text-decoration: line-through\">Strikethrough paragraph with "
            "<span style=\"text-decoration: underline red\">underlined red span</span> inside</p>"

            # Override behavior with individual properties
            "<p style=\"text-decoration-line: underline; text-decoration-color: blue\">Blue underlined paragraph with "
            "<span style=\"text-decoration-line: line-through;\">strikethrough span</span> inside</p>"

            # Check if equal - shorthand vs individual properties
            "<p style=\"text-decoration-line: underline; text-decoration-color: blue; text-decoration-style: wavy\">Blue underlined paragraph</p>"
            "<p style=\"text-decoration: underline blue wavy\">Blue underlined paragraph</p>"
        )

        self.parser.add_html_to_document(text_decoration_html_example, self.document)
        document = self.parser.parse_html_string(text_decoration_html_example)

        # --------------------------------------------------------------------
        # 1) Standalone spans
        # --------------------------------------------------------------------
        standalone = document.paragraphs[0].runs
        assert len(standalone) == 4

        span1 = standalone[0]
        assert span1.text == "underlined span (red)"
        assert span1.font.underline == WD_UNDERLINE.SINGLE
        assert self.get_underline_color(span1) == self.hexcolor("red")

        span2 = standalone[1]
        assert span2.text == "no decoration span (rgb(0, 0, 0))"
        assert span2.font.underline is False
        assert span2.font.strike is False
        assert self.get_underline_color(span2) == self.hexcolor("rgb(0,0,0)")

        span3 = standalone[2]
        assert span3.text == "strikethrough span (gray) (not supported)"
        assert span3.font.strike is True
        assert self.get_underline_color(span3) == self.hexcolor("gray")

        span4 = standalone[3]
        assert span4.text == "underline+line-through span (orange) (should be strike)"
        assert span4.font.strike is True
        assert span4.font.underline is False
        assert self.get_underline_color(span4) == self.hexcolor("orange")

        # --------------------------------------------------------------------
        # 2) Spans inside paragraphs (wavy, dotted, strikethrough)
        # --------------------------------------------------------------------
        p1 = document.paragraphs[1]
        r = p1.runs[1]
        assert r.text == "wavy underlined span (blue)"
        assert r.font.underline == WD_UNDERLINE.WAVY
        assert self.get_underline_color(r) == self.hexcolor("blue")

        p2 = document.paragraphs[2]
        r = p2.runs[1]
        assert r.text == "dotted underlined span (purple)"
        assert r.font.underline == WD_UNDERLINE.DOTTED
        assert self.get_underline_color(r) == self.hexcolor("purple")

        p3 = document.paragraphs[3]
        r = p3.runs[1]
        assert r.text == "strikethrough span (red)"
        assert r.font.strike is True
        assert self.get_underline_color(r) == self.hexcolor("red")

        # --------------------------------------------------------------------
        # 3) Paragraph with multiple spans
        # --------------------------------------------------------------------
        p4 = document.paragraphs[4]
        runs = p4.runs

        multiple_span1 = runs[1]
        assert multiple_span1.text == "underlined"
        assert multiple_span1.font.underline == WD_UNDERLINE.SINGLE
        assert multiple_span1.font.strike is False
        assert self.get_underline_color(multiple_span1) == self.hexcolor("green")

        multiple_span2 = runs[3]
        assert multiple_span2.text == "strikethrough"
        assert multiple_span2.font.underline is False
        assert multiple_span2.font.strike is True
        assert self.get_underline_color(multiple_span2) == self.hexcolor("blue")

        multiple_span3 = runs[5]
        assert multiple_span3.text == "dashed underline"
        assert multiple_span3.font.underline == WD_UNDERLINE.DASH
        assert multiple_span3.font.strike is False
        assert self.get_underline_color(multiple_span3) == self.hexcolor("orange")

        # --------------------------------------------------------------------
        # 4) Underlined paragraph + span with none
        # --------------------------------------------------------------------
        p5 = document.paragraphs[5]
        assert p5.runs[0].font.underline == WD_UNDERLINE.SINGLE

        span_none = p5.runs[1]
        assert span_none.font.underline is False
        assert span_none.font.strike is False

        # --------------------------------------------------------------------
        # 5) Line-through paragraph + underlined red span
        # --------------------------------------------------------------------
        p6 = document.paragraphs[6]
        assert p6.runs[0].font.strike is True

        span = p6.runs[1]
        assert span.font.underline is True
        assert self.get_underline_color(span) == self.hexcolor("red")

        # --------------------------------------------------------------------
        # 6) Individual properties
        # --------------------------------------------------------------------
        p7 = document.paragraphs[7]
        r = p7.runs[0]
        assert r.font.underline is True
        assert self.get_underline_color(r) == self.hexcolor("blue")

        span = p7.runs[1]
        assert span.font.strike is True

        # --------------------------------------------------------------------
        # 7) Shorthand vs individual comparison
        # --------------------------------------------------------------------
        p8 = document.paragraphs[8]
        p9 = document.paragraphs[9]

        assert p8.runs[0].font.underline == WD_UNDERLINE.WAVY
        assert p9.runs[0].font.underline == WD_UNDERLINE.WAVY

        assert p8.text == "Blue underlined paragraph"
        assert p9.text == "Blue underlined paragraph"

        assert self.get_underline_color(p8.runs[0]) == self.hexcolor("blue")
        assert self.get_underline_color(p9.runs[0]) == self.hexcolor("blue")

    def test_text_decoration_paragraph(self):
        self.document.add_heading('Test: text-decoration on <p>', level=1)
        text_decoration_html_example = (
            "<p style=\"text-decoration: underline red\">underlined text (red)</p>"
            "<p style=\"text-decoration: none rgb(0,0,0)\">no decoration text (rgb(0, 0, 0))</p>"
            "<p style=\"text-decoration: line-through gray\">strikethrough text (gray) (color not supported)</p>"
            "<p style=\"text-decoration: underline line-through orange\">underline+line-through (orange)\
                (should be strike)</p>"
            "<p style=\"text-decoration: underline wavy blue\">wavy underline (blue)</p>"
            "<p style=\"text-decoration: underline dotted rgb(0, 128, 0)\">dotted underline (rgb(0, 128, 0))</p>"
            "<p style=\"text-decoration: underline dotted rgb(0, 255, 0)\">dotted underline (rgb(0, 255, 0))</p>"
            "<p style=\"text-decoration: underline dashed purple\">dashed underline (purple)</p>"
            "<p style=\"text-decoration: underline double rgb(255, 69, 0)\">double underline (rgb(255, 69, 0))</p>"
            "<p style=\"text-decoration: overline hotpink\">overline text (hotpink) (not supported)</p>"
            "<p style=\"text-decoration: blink hotpink\">blink text (hotpink) (not supported)</p>"
        )

        self.parser.add_html_to_document(text_decoration_html_example, self.document)
        with self.assertLogs(level='WARNING') as log:
            document = self.parser.parse_html_string(text_decoration_html_example)

        underline_states = []
        underline_colors = []
        strike_states = []

        for p in document.paragraphs:
            run = p.runs[0]

            # Check underline
            underline = run.font.underline
            if underline is None:
                underline_states.append(None)
                underline_colors.append(None)
            elif underline is True:
                underline_states.append(True)
                underline_colors.append(self.get_underline_color(run))
            elif underline is False:
                underline_states.append(False)
                underline_colors.append(None)
            else:
                underline_states.append(underline)
                underline_colors.append(self.get_underline_color(run))

            # Check strike-through
            strike = run.font.strike
            if strike is None:
                strike_states.append(None)
            elif strike is True:
                strike_states.append(True)
            elif strike is False:
                strike_states.append(False)
            else:
                strike_states.append(strike)

        expected_underline_states = [
            True,   # underline (default single) - explicitly True
            False,  # none - explicitly False for both underline and strike
            False,  # line-through - explicitly False for underline when strike is True
            False,  # underline + line-through - line-through wins, underline explicitly False
            WD_UNDERLINE.WAVY,      # wavy underline - explicitly set to wavy
            WD_UNDERLINE.DOTTED,    # dotted underline - explicitly set to dotted
            WD_UNDERLINE.DOTTED,    # dotted underline - explicitly set to dotted
            WD_UNDERLINE.DASH,      # dashed underline - explicitly set to dash
            WD_UNDERLINE.DOUBLE,    # double underline - explicitly set to double
            None,  # overline (not supported) - remains None/unchanged
            None,  # blink (not supported) - remains None/unchanged
        ]

        expected_underline_colors = [
            self.hexcolor("red"),             # underline red
            None,                             # none rgb(0,0,0)
            None,                             # line-through gray (strike only, but color captured)
            None,                             # underline + line-through (color should be orange)
            self.hexcolor("blue"),            # wavy underline blue
            self.hexcolor("rgb(0,128,0)"),    # dotted underline rgb(0,128,0)
            self.hexcolor("rgb(0,255,0)"),    # dotted underline rgb(0,255,0)
            self.hexcolor("purple"),          # dashed underline purple
            self.hexcolor("rgb(255,69,0)"),   # double underline rgb(255,69,0)
            None,                             # overline hotpink (unsupported → underline None, but color still parsed)
            None,                             # blink hotpink (unsupported)
        ]

        expected_strike_states = [
            False,  # underline only - explicitly False for strike when underline is True
            False,  # none - explicitly False for both underline and strike
            True,   # line-through - explicitly True
            True,   # underline + line-through - line-through wins, strike explicitly True
            False,  # wavy underline only - explicitly False for strike when underline is set
            False,  # dotted underline only - explicitly False for strike when underline is set
            False,  # dotted underline only - explicitly False for strike when underline is set
            False,  # dashed underline only - explicitly False for strike when underline is set
            False,  # double underline only - explicitly False for strike when underline is set
            None,   # overline (not supported) - remains None/unchanged
            None,   # blink (not supported) - remains None/unchanged
        ]

        # Test that the underline states, colors, and strike states are correct
        self.assertEqual(underline_states, expected_underline_states)
        self.assertEqual(underline_colors, expected_underline_colors)
        self.assertEqual(strike_states, expected_strike_states)

        # Test that the correct warnings are logged
        self.assertEqual(len(log.records), 4)
        self.assertIn('Word does not support colored strike-through. Color \'gray\' will be ignored for line-through.', log.output[0])
        self.assertIn('Word does not support colored strike-through. Color \'orange\' will be ignored for line-through.', log.output[1])
        self.assertIn('Blink or overline not supported.', log.output[2])
        self.assertIn('Blink or overline not supported.', log.output[3])

    def test_first_line_paragraph(self):
        self.document.add_heading('Test text-indent on <p> tags', level=1)
        self.parser.add_html_to_document(self.paragraph_first_line_indent, self.document)
        document = self.parser.parse_html_string(self.paragraph_first_line_indent)

        indent_values = []

        for p in document.paragraphs:
            indent_pt = p.paragraph_format.first_line_indent
            if indent_pt is not None:
                indent_values.append(indent_pt)

        expected_values = [
            1080000,    # 3cm
            254000,     # 20pt
            381000,     # 40px
            1260000,    # 35mm
            None,       # Word does not support negative values here
        ]

        for actual, expected in zip(indent_values, expected_values):
            self.assertAlmostEqual(actual, expected, delta=634)

    def test_color_paragraph(self):
        self.document.add_heading('Test: color on p tags', level=1)
        color_html_example = (
            "<p style=\"color: red\">red text</p>"
            "<p style=\"color: #00ff00\">green hex text</p>"
            "<p style=\"color: rgb(0, 0, 255)\">blue rgb text</p>"
            "<p style=\"color: inherit\">inherit color text</p>"
            "<p style=\"color: transparent\">transparent color text</p>"
            "<p style=\"color: currentcolor\">current color text</p>"
            "<p style=\"color: #ff0000; font-size: 14pt\">red with other styles</p>"
            "<p>default text</p>"
        )

        self.parser.add_html_to_document(color_html_example, self.document)

        document = self.parser.parse_html_string(color_html_example)

        color_states = []
        for p in document.paragraphs:
            if p.runs and p.runs[0].font.color:
                color_rgb = p.runs[0].font.color.rgb
                if color_rgb:
                    color_states.append((color_rgb[0], color_rgb[1], color_rgb[2]))
                else:
                    color_states.append(None)
            else:
                color_states.append(None)

        expected_colors = [
            (255, 0, 0),  # red
            (0, 255, 0),  # #00ff00 (green)
            (0, 0, 255),  # rgb(0, 0, 255) (blue)
            None,  # inherit (should not apply color)
            None,  # transparent (should not apply color)
            None,  # currentcolor (should not apply color)
            (255, 0, 0),  # #ff0000 (red) with other styles
            None,  # default text
        ]

        self.assertEqual(color_states, expected_colors)

    def test_line_height_paragraph(self):
        self.document.add_heading('Test: line-height on <p>', level=1)
        self.parser.add_html_to_document(self.paragraph_line_height, self.document)
        document = self.parser.parse_html_string(self.paragraph_line_height)

        line_heights = []
        line_rules = []

        for p in document.paragraphs:
            line_spacing = p.paragraph_format.line_spacing
            line_rule = p.paragraph_format.line_spacing_rule
            line_heights.append(str(line_spacing) if line_spacing is not None else 'None')
            line_rules.append(str(line_rule) if line_rule is not None else 'None')

        expected_line_heights = [
            '1.0',
            '1.15',
            '1.5',
            '2.0',
            '190500',   # line-height: 20px
            '182880',   # line-height: 1.2em
            '228600',   # line-height: 1.5em
            '304800',   # line-height: 2em
            '182880',   # line-height: 1.2rem
            '228600',   # line-height: 1.5rem
            '304800',   # line-height: 2rem
            '1.5',      # line-height: 150%
            '2.0',      # line-height: 200%
        ]

        self.assertEqual(line_heights, expected_line_heights,
                         f"Line heights don't match expected values. Got {line_heights}, expected {expected_line_heights}")

    def test_margins_paragraph(self):
        margins_html_example = (
            "<p style=\"margin-left: auto; margin-right: auto\">centered paragraph</p>"
            "<p style=\"margin-left: 20px\">left margin 20px</p>"
            "<p style=\"margin-right: 1.5cm\">right margin 1.5cm</p>"
            "<p style=\"margin-left: 1cm\">left margin 1cm</p>"
            "<p style=\"margin-left: 10px; margin-right: 15px\">both margins set</p>"
            "<p style=\"margin-left: auto\">only left auto</p>"
            "<p style=\"margin-right: auto\">only right auto</p>"
            "<p style=\"margin-left: 0px; margin-right: 0px\">zero margins</p>"
            "<p style=\"margin-left: 2in\">left margin 2in</p>"
        )

        self.document.add_heading('Test margins on <p>', level=1)
        self.parser.add_html_to_document(margins_html_example, self.document)
        document = self.parser.parse_html_string(margins_html_example)

        expected_margins = [
            # Paragraph 1: "centered paragraph" - auto margins (None values)
            {'left': None, 'right': None},
            # Paragraph 2: "left margin 20px" - 20px = 20 * 9525 = 190500 EMU
            {'left': 190500, 'right': None},
            # Paragraph 3: "right margin 1.5cm" - 1.5cm = 1.5 * 360000 = 540000 EMU
            {'left': None, 'right': 540000},
            # Paragraph 4: "left margin 1cm" - 1cm = 360000 EMU
            {'left': 360000, 'right': None},
            # Paragraph 5: "both margins set" - 10px=95250 EMU, 15px=142875 EMU
            {'left': 95250, 'right': 142875},
            # Paragraph 6: "only left auto" - auto margin
            {'left': None, 'right': None},
            # Paragraph 7: "only right auto" - auto margin
            {'left': None, 'right': None},
            # Paragraph 8: "zero margins" - 0px = 0 EMU
            {'left': 0, 'right': 0},
            # Paragraph 9: "left margin 2in" - 2in = 2 * 914400 = 1828800 EMU
            {'left': 1828800, 'right': None},
        ]

        self.assertEqual(len(document.paragraphs), len(expected_margins))

        for i, paragraph in enumerate(document.paragraphs):
            expected = expected_margins[i]
            actual_left = paragraph.paragraph_format.left_indent
            actual_right = paragraph.paragraph_format.right_indent

            # Check left margin
            if expected['left'] is None:
                self.assertIsNone(actual_left, f"Paragraph {i} left margin should be None")
            else:
                self.assertIsNotNone(actual_left, f"Paragraph {i} left margin should not be None")
                self.assertTrue(abs(actual_left - expected['left']) <= 634,
                                f"Paragraph {i} left margin: expected {expected['left']} EMU, got {actual_left} EMU")

            # Check right margin
            if expected['right'] is None:
                self.assertIsNone(actual_right, f"Paragraph {i} right margin should be None")
            else:
                self.assertIsNotNone(actual_right, f"Paragraph {i} right margin should not be None")
                self.assertTrue(abs(actual_right - expected['right']) <= 634,
                                f"Paragraph {i} right margin: expected {expected['right']} EMU, got {actual_right} EMU")

    def test_background_color_styles(self):
        self.document.add_heading('Test background color on <p>, multiple cases', level=1)
        html_example2 = """
        <p style="background-color: lightblue;">
            Start of paragraph
            <span style="background-color: yellow;">First yellow span</span>
            middle text
            <span style="background-color: red; color: white;">Red span with white text</span>
            end of paragraph
            <span style="background-color: purple; color: white;">
                Purple span with
                <span style="background-color: orange;">nested orange</span>
                inside
            </span>
        </p>
        """
        self.parser.add_html_to_document(html_example2, self.document)

        html_example3 = """
        <p style="background-color: #f0f0f0;">
            Base paragraph background
            <span style="background-color: #ffcccc; font-weight: bold;">Bold pink span</span>
            regular text
            <span style="background-color: #ccffcc;">
                Green span with
                <span style="background-color: #ccccff; font-style: italic;">italic blue nested</span>
                and more green
            </span>
            <span style="background-color: #ffffcc;">
                Yellow span with
                <span style="background-color: #ffccff;">pink nested</span>
                and
                <span style="background-color: #ccffff; text-decoration: underline;">cyan underlined</span>
            </span>
        </p>
        """
        self.parser.add_html_to_document(html_example3, self.document)

        html_example4 = """
        <p style="background-color: white;">
            White paragraph
            <span style="background-color: yellow;">Yellow span</span>
            <span style="background-color: transparent;">Transparent span</span>
        </p>
        """
        self.parser.add_html_to_document(html_example4, self.document)

        html_example5 = """
        <p style="background-color: rgb(200, 200, 255);">
            RGB color background
            <span style="background-color: #ff0000;">Hex red</span>
            <span style="background-color: inherit;">Inherit background</span>
            <span style="background-color: initial;">Initial background</span>
        </p>
        """
        self.parser.add_html_to_document(html_example5, self.document)

        html_example6 = """
        <p style="background-color: #e8e8e8;">
            Level 0
            <span style="background-color: #d0d0ff;">
                Level 1
                <span style="background-color: #ffd0d0;">
                    Level 2
                    <span style="background-color: #d0ffd0;">
                        Level 3
                        <span style="background-color: #ffffd0;">
                            Level 4
                            <span style="background-color: #ffd0ff;">
                                Level 5
                            </span>
                        </span>
                    </span>
                </span>
            </span>
            Back to level 0
        </p>
        """
        self.parser.add_html_to_document(html_example6, self.document)

        html_example7 = """
        <p style="background-color: #f5f5f5; padding: 10px;">
            Paragraph with padding
            <span style="background-color: yellow; font-size: 14pt; font-weight: bold;">Styled span</span>
            <span style="background-color: lightgreen; text-decoration: underline; color: darkblue;">Underlined green</span>
            <span style="background-color: pink; font-family: 'Arial'; font-style: italic;">Italic Arial pink</span>
        </p>
        """
        self.parser.add_html_to_document(html_example7, self.document)

        html_example8 = """
        <p>
            Normal paragraph
            <span style="background-color: #ffeb3b;">Highlighted text</span>
            normal text
            <span style="background-color: #4caf50; color: white;">Green highlight</span>
            more normal text
        </p>
        """
        self.parser.add_html_to_document(html_example8, self.document)

        html_example9 = """
        <p style="background-color: #fffacd;">
            Light yellow background entire paragraph
        </p>
        <p>
            No background
            <span style="background-color: #ffcdd2;">Light red span only</span>
        </p>
        <p style="background-color: #e1f5fe;">
            Light blue background
            <span style="background-color: #c8e6c9;">Light green span</span>
            <span style="background-color: #ffecb3;">Light orange span</span>
        </p>
        """
        self.parser.add_html_to_document(html_example9, self.document)

    def test_headers_with_css(self):
        self.document.add_heading('Test: headers with css', level=1)
        self.parser.add_html_to_document(self.css_properties_header, self.document)

        document = self.parser.parse_html_string(self.css_properties_header)

        # Test H1 - Large and Centered
        h1 = document.paragraphs[0]
        assert h1.style.name.startswith('Heading 1')
        assert str(h1.runs[0].font.color.rgb) == '2C3E50'
        assert h1.runs[0].font.bold is True
        assert h1.runs[0].font.size == 342900
        assert h1.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert h1.runs[0].text == 'MAIN HEADING H1 - LARGE AND CENTERED'  # uppercase due to text-transform

        # Test H2 - Underlined with Background (no span in this one)
        h2 = document.paragraphs[1]
        assert h2.style.name.startswith('Heading 2')
        assert str(h2.runs[0].font.color.rgb) == '34495E'
        assert h2.runs[0].font.underline is True
        assert h2.runs[0].font.name == 'Arial'
        assert h2.runs[0].font.size == 266700

        # Test H3 - Italic and Right Aligned
        h3 = document.paragraphs[2]
        assert h3.style.name.startswith('Heading 3')
        assert str(h3.runs[0].font.color.rgb) == '7F8C8D'
        assert h3.runs[0].font.italic is True
        assert h3.runs[0].font.size == 209550
        assert h3.alignment == WD_ALIGN_PARAGRAPH.RIGHT

        # Test H4 - Normal Weight and Capitalized
        h4 = document.paragraphs[3]
        assert h4.style.name.startswith('Heading 4')
        assert str(h4.runs[0].font.color.rgb) == '95A5A6'
        assert h4.runs[0].font.bold is False  # font-weight: normal
        assert h4.runs[0].font.name == 'Georgia'
        assert h4.runs[0].font.size == 171450
        assert h4.runs[0].text == 'Quaternary Heading H4 - Normal Weight And Capitalized'  # capitalized

        # Test H1 with Complex Text Decoration and Span
        h1_complex = document.paragraphs[4]
        assert h1_complex.runs[0].font.strike is True  # line-through
        assert str(h1_complex.runs[0].font.color.rgb) == '8E44AD'
        assert h1_complex.runs[0].font.size == 381000

        # Test span in complex H1
        assert len(h1_complex.runs) >= 2
        span_in_h1 = h1_complex.runs[1]
        assert span_in_h1.font.underline is True  # underline in span
        assert str(span_in_h1.font.color.rgb) == '2980B9'

        # Test H3 with Light Weight and Span
        h3_light = document.paragraphs[5]
        assert h3_light.runs[0].font.bold is False  # font-weight: 100
        assert str(h3_light.runs[0].font.color.rgb) == 'D35400'
        assert h3_light.runs[0].font.size == 190500

        # Test bold span in light H3
        assert len(h3_light.runs) >= 2
        bold_span = h3_light.runs[1]
        assert bold_span.font.bold is True  # font-weight: 900

        # Test H3 with Text Transform
        h3_transform = document.paragraphs[6]
        assert h3_transform.runs[0].text == 'h3 forced to lowercase with text-transform '
        assert len(h3_transform.runs) >= 2
        uppercase_span = h3_transform.runs[1]
        assert uppercase_span.text == 'SPAN FORCED TO UPPERCASE'

        # Test H4 with Serif Font
        h4_serif = document.paragraphs[7]
        assert h4_serif.runs[0].font.name == 'Times New Roman'
        assert str(h4_serif.runs[0].font.color.rgb) == '7D3C98'
        assert h4_serif.alignment == WD_ALIGN_PARAGRAPH.CENTER

        # Test H1 with Auto Margins and Background
        h1_centered = document.paragraphs[8]
        assert h1_centered.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert str(h1_centered.runs[0].font.color.rgb) == 'FFFFFF'

        # Test H2 with Lighter Weight and Span
        h2_lighter = document.paragraphs[9]
        assert h2_lighter.runs[0].font.bold is False  # lighter weight
        assert h2_lighter.runs[0].font.underline == WD_UNDERLINE.DOTTED
        assert h2_lighter.runs[0].font.size == 228600

        # Test bolder span
        assert len(h2_lighter.runs) >= 2
        bolder_span = h2_lighter.runs[1]
        assert bolder_span.font.bold is True  # bolder

        # Test H3 with RGB Colors and Span
        h3_rgb = document.paragraphs[10]
        assert str(h3_rgb.runs[0].font.color.rgb) == '3498DB'  # rgb(52, 152, 219)
        assert h3_rgb.runs[0].font.size == 177800

        # Test RGB span
        assert len(h3_rgb.runs) >= 2
        rgb_span = h3_rgb.runs[1]
        assert str(rgb_span.font.color.rgb) == 'E74C3C'  # rgb(231, 76, 60)

        # Test H4 with Strike-through and Span
        h4_strike = document.paragraphs[11]
        assert h4_strike.runs[0].font.strike is True
        assert h4_strike.runs[0].font.bold is False  # font-weight: 300

        # Test span without strike-through
        assert len(h4_strike.runs) >= 2
        no_strike_span = h4_strike.runs[1]
        assert no_strike_span.font.strike is False
        assert str(no_strike_span.font.color.rgb) == 'E74C3C'

        # Test H3 with Unsupported Transform and Span
        h3_unsupported = document.paragraphs[12]
        assert str(h3_unsupported.runs[0].font.color.rgb) == 'F39C12'
        assert h3_unsupported.runs[0].font.size == 196850

        # Test supported transform in span
        assert len(h3_unsupported.runs) >= 2
        supported_span = h3_unsupported.runs[1]
        assert supported_span.text == 'Supported Transform In Span'  # capitalize

        # Test H4 with Reset Styles and Span
        h4_reset = document.paragraphs[13]
        assert h4_reset.runs[0].font.bold is True  # font-weight: 700
        assert h4_reset.runs[0].font.italic is False  # font-style: normal
        assert h4_reset.runs[0].font.underline is False  # text-decoration: none

        # Test styled span
        assert len(h4_reset.runs) >= 2
        styled_span = h4_reset.runs[1]
        assert styled_span.font.bold is False  # font-weight: 400
        assert styled_span.font.italic is True
        assert styled_span.font.underline is True

        # Test H1 with Text Color and Span
        h1_transparent = document.paragraphs[14]
        assert h1_transparent.runs[0].font.size == 361950
        visible_span = h1_transparent.runs[1]
        assert str(visible_span.font.color.rgb) == 'ECF0F1'

        # Test H3 with All Three Decorations and Span
        h3_all_decorations = document.paragraphs[15]
        assert h3_all_decorations.runs[0].font.strike is True
        assert h3_all_decorations.runs[0].font.underline is False

        # Test span with single decoration
        assert len(h3_all_decorations.runs) >= 2
        single_decoration_span = h3_all_decorations.runs[1]
        assert single_decoration_span.font.underline is True

        # Test H2 with Middle Weight and Span
        h2_middle = document.paragraphs[16]
        assert h2_middle.runs[0].font.bold is False

        # Test darker span
        assert len(h2_middle.runs) >= 2
        darker_span = h2_middle.runs[1]
        assert darker_span.font.bold is False

        # Test H4 with Style and Span
        h4_style = document.paragraphs[17]
        assert h4_style.runs[0].font.underline is WD_UNDERLINE.WAVY
        assert h4_style.runs[1].font.underline is WD_UNDERLINE.DOUBLE

    def test_color_by_name(self):
        color_html_example = (
            "<p><span style=\"color:red\">paragraph red</span></p>"
            "<p><span style=\"color: yellow\">paragraph yellow</span></p>"
            "<p><span style=\"color: blue !important\">paragraph blue</span></p>"
            "<p><span style=\"color: green!important\">paragraph green</span></p>"
            "<p><span style=\"color: darkgray!IMPORTANT\">paragraph darkgray</span></p>"
            "<p><span style=\"color: MAGENTA !IMPORTANT\">paragraph magenta</span></p>"
            "<p><span style=\"color: invalidcolor\">paragraph has default black because of invalid color name</span></p>"
        )

        self.document.add_heading(
            'Test: Color by name',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(color_html_example, self.document)

        document = self.parser.parse_html_string(color_html_example)
        colors = [str(p.runs[0].font.color.rgb) for p in document.paragraphs]

        assert 'FF0000' in colors # Red
        assert 'FFFF00' in colors # Yellow
        assert '0000FF' in colors # Blue
        assert '008000' in colors # Green
        assert 'A9A9A9' in colors # Darkgray
        assert '000000' in colors # Black
        assert 'FF00FF' in colors # Magenta

    def test_table_cell_border_properties(self):
        """Validates that all table cells have the expected border size, style, and color."""

        self.document.add_heading(
            'Test: Table Cell Border Properties',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(self.table3_html, self.document)
        document = self.parser.parse_html_string(self.table3_html)

        # Define expected border properties
        expected_colors = [
            {
                "top": {"color": "D95B48", "style": "single", "size": "1.0"},
                "bottom": {"color": "D95B48", "style": "single", "size": "1.0"},
                "left": {"color": "FF0000", "style": "single", "size": "1.0"},
                "right": {"color": "8B0000", "style": "single", "size": "1.0"}
            },
            {
                "top": {"color": "FAC32A", "style": "single", "size": "1.0"},
                "bottom": {"color": "FAC32A", "style": "single", "size": "1.125"},
                "left": {"color": "none", "style": "none", "size": "none"},
                "right": {"color": "FAC32A", "style": "single", "size": "12.0"}
            },
            {
                "top": {"color": "30E667", "style": "none", "size": "5.67"},
                "bottom": {"color": "30E667", "style": "single", "size": "5.67"},
                "left": {"color": "30E667", "style": "single", "size": "5.67"},
                "right": {"color": "30E667", "style": "single", "size": "5.67"}
            },
            {
                "top": {"color": "none", "style": "none", "size": "none"},
                "bottom": {"color": "D948CF", "style": "single", "size": "1.5"},
                "left": {"color": "none", "style": "none", "size": "none"},
                "right": {"color": "D948CF", "style": "single", "size": "5.67"}
            },
            {
                "top": {"color": "EAAAA7", "style": "single", "size": "1.1"},
                "bottom": {"color": "EAAAA7", "style": "single", "size": "1.1"},
                "left": {"color": "EAAAA7", "style": "single", "size": "1.1"},
                "right": {"color": "EAAAA7", "style": "single", "size": "1.1"}
            },
            {
                "top": {"color": "none", "style": "none", "size": "none"},
                "bottom": {"color": "ACC4AA", "style": "dashed", "size": "7.2"},
                "left": {"color": "none", "style": "none", "size": "none"},
                "right": {"color": "ACC4AA", "style": "dotted", "size": "4.8"}
            }
        ]

        # Validate border properties for each cell
        cell_idx = 0
        for row_idx, row in enumerate(document.tables[0].rows):
            for column_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))

                # Extract border properties
                border_sides = {
                    'top': tcBorders.find(qn('w:top')) if tcBorders is not None else None,
                    'bottom': tcBorders.find(qn('w:bottom')) if tcBorders is not None else None,
                    'left': tcBorders.find(qn('w:left')) if tcBorders is not None else None,
                    'right': tcBorders.find(qn('w:right')) if tcBorders is not None else None,
                }

                for side, border in border_sides.items():
                    if border is not None:
                        color = border.get(qn('w:color'), "").upper()  # Ensure uppercase and no #
                        size = border.get(qn('w:sz'))
                        style = border.get(qn('w:val'))
                    else:
                        color, size, style = "none", "none", "none"

                    # Convert size from eighths of a point to points
                    size_in_pt = str(round(float(size) / 8, 3)) if size and size != "none" else "none"

                    # Get expected properties for the current cell and side
                    expected_properties = expected_colors[cell_idx][side]

                    # Assertions
                    assert color == expected_properties["color"], (
                        f"Color mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['color']}, got {color}"
                    )
                    assert size_in_pt == expected_properties["size"], (
                        f"Size mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['size']}, got {size_in_pt}"
                    )
                    assert style == expected_properties["style"], (
                        f"Style mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['style']}, got {style}"
                    )

                cell_idx += 1

    def test_table_cell_background_color(self):
        """Validates that all table cells have the expected background color."""

        self.document.add_heading(
            'Test: Table Cell Background Color',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(self.table3_html, self.document)
        document = self.parser.parse_html_string(self.table3_html)

        # Define expected background colors for each cell
        expected_background_colors = [
            "3749EF", # Row 1 Column 1
            "33b32e", # Row 1 Column 2
            "BFBFBF", # Row 2 Column 1
            "2eaab3", # Row 2 Column 2
            "99fffa", # Row 3 Column 1
            "2eaab3"  # Row 3 Column 2
        ]

        # Validate background colors for each cell
        cell_idx = 0
        for row_idx, row in enumerate(document.tables[0].rows):
            for column_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

                # Get the background color (shading) if it exists
                shading = tcPr.find(qn('w:shd'))
                if shading is not None:
                    background_color = shading.get(qn('w:fill'), "").upper()  # Ensure uppercase and no #
                else:
                    background_color = "None"

                # Get expected background color for the current cell
                expected_color = expected_background_colors[cell_idx].upper()
                cell_idx += 1

                assert background_color == expected_color, (
                    f"Background color mismatch for row {row_idx} column {column_idx}: "
                    f"expected {expected_color}, got {background_color}"
                )

    def test_table_cell_dimensions(self):
        """Validates that all table cells have the expected width and height."""

        self.document.add_heading(
            'Test: Table Cell Dimensions',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(self.table3_html, self.document)
        document = self.parser.parse_html_string(self.table3_html)

        # Define expected dimensions for each cell
        expected_dimensions = [
            # First row
            [
                {
                    "width": "258.35px",  # Width for the first cell
                    "height": "23.75pt"   # Height for the first cell
                },
                {
                    "width": "222.2pt",   # Width for the second cell
                    "height": "23.75pt"   # Height for the second cell
                }
            ],
            # Second row
            [
                {
                    "width": "258.35in",  # Width for the first cell
                    "height": "15.5pt"    # Height for the first cell
                },
                {
                    "width": "6cm",       # Width for the second cell
                    "height": "15.5pt"    # Height for the second cell
                }
            ],
            # Third row
            [
                {
                    "width": "258.35pt",  # Width for the first cell
                    "height": "2rem"      # Height for the first cell
                },
                {
                    "width": "6cm",       # Width for the second cell
                    "height": "2em"       # Height for the second cell
                }
            ]
        ]

        # Validate dimensions for each cell
        for row_idx, row in enumerate(document.tables[0].rows):
            for cell_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                docx_cell = document.tables[0].cell(row_idx, cell_idx)

                # Convert width from EMUs to px
                cell_width_px = round((docx_cell.width / 914400) * 96, 2)  # 1 EMU = 1/914400 inch, 1 inch = 96px
                # Get expected width and convert it to points using unit_converter
                expected_width = expected_dimensions[row_idx][cell_idx]["width"]
                expected_width_px = unit_converter(expected_width, "px")

                # Convert height from EMUs to px
                cell_height_px = round((row.height / 914400) * 96, 2)  # 1 EMU = 1/914400 inch, 1 inch = 96px
                # Get expected height and convert it to points
                expected_height = expected_dimensions[row_idx][cell_idx]["height"]
                expected_height_px = unit_converter(expected_height, "px")

                assert round(abs(cell_width_px - expected_width_px), 2) <= 0.03, (
                    f"Width mismatch for cell ({row_idx}, {cell_idx}): "
                    f"expected {expected_width_px}px, got {cell_width_px}px"
                )
                assert round(abs(cell_height_px - expected_height_px), 2) <= 0.03, (
                    f"Height mismatch for cell ({row_idx}, {cell_idx}): "
                    f"expected {expected_height_px}px, got {cell_height_px}px"
                )

    def test_border_with_keywords(self):
        self.document.add_heading("Test: Cells with size keywords", level=1)

        size_keywords_html_example = """
        <table>
        <tr>
            <td style="border: thin solid currentcolor">Hello World</td>
            <td style="border: medium dashed orange">Thanks World</td>
        </tr>
        <tr>
            <td style="border: thick double red">Goodbye World</td>
            <td style="border: none">Bye World</td>
        </tr>
        </table>
        """
        self.parser.add_html_to_document(size_keywords_html_example, self.document)
        document = self.parser.parse_html_string(size_keywords_html_example)

        expected_sizes = ["0.75", "2.25", "3.75", "none"]

        # Validate border properties for each cell
        cell_idx = 0
        for row_idx, row in enumerate(document.tables[0].rows):
            for column_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))

                # Extract border properties
                border_sides = {
                    'top': tcBorders.find(qn('w:top')) if tcBorders is not None else None,
                    'bottom': tcBorders.find(qn('w:bottom')) if tcBorders is not None else None,
                    'left': tcBorders.find(qn('w:left')) if tcBorders is not None else None,
                    'right': tcBorders.find(qn('w:right')) if tcBorders is not None else None,
                }

                for side, border in border_sides.items():
                    size = border.get(qn('w:sz')) if border is not None else "none"

                    # Convert size from eighths of a point to points
                    size_in_pt = str(float(size) / 8) if size and size != "none" else "none"

                    assert size_in_pt == expected_sizes[cell_idx], (
                        f"Size mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_sizes[cell_idx]}, got {size_in_pt}"
                    )

                cell_idx += 1

    def test_border_style_with_diff_formats(self):
        self.document.add_heading("Test: Cells border style with different formats", level=1)

        size_keywords_html_example = """
        <table>
            <tr>
                <td style="border-top: lightblue; border-left: medium none currentcolor; border-right: solid; border-bottom: none">Cell 1</td>
                <td style="border-top: 5px; border-left:  1px  solid  black ; border-right: thin #736; border-bottom:  ;">Cell 2</td>
                <td style="border-top: orange solid; border-bottom: magenta solid thick; border-left: medium dashed currentcolor;">Cell 3</td>
            </tr>
        </table>
        """
        self.parser.add_html_to_document(size_keywords_html_example, self.document)
        document = self.parser.parse_html_string(size_keywords_html_example)

        expected_sides = [
            {
                "top": {"color": "ADD8E6", "style": "single", "size": "1.0"},
                "bottom": {"color": "none", "style": "none", "size": "none"},
                "left": {"color": "000000", "style": "none", "size": "2.25"},
                "right": {"color": "000000", "style": "single", "size": "1.0"}
            },
            {
                "top": {"color": "000000", "style": "single", "size": "3.75"},
                "bottom": {"color": "none", "style": "none", "size": "none"},
                "left": {"color": "000000", "style": "single", "size": "0.75"},
                "right": {"color": "773366", "style": "single", "size": "0.75"}
            },
            {
                "top": {"color": "FFA500", "style": "single", "size": "1.0"},
                "bottom": {"color": "FF00FF", "style": "single", "size": "3.75"},
                "left": {"color": "000000", "style": "dashed", "size": "2.25"},
                "right": {"color": "none", "style": "none", "size": "none"}
            }
        ]

        cell_idx = 0
        for row_idx, row in enumerate(document.tables[0].rows):
            for column_idx, cell in enumerate(row.cells):
                # Get the table cell element and properties
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))

                # Extract border properties
                border_sides = {
                    'top': tcBorders.find(qn('w:top')) if tcBorders is not None else None,
                    'bottom': tcBorders.find(qn('w:bottom')) if tcBorders is not None else None,
                    'left': tcBorders.find(qn('w:left')) if tcBorders is not None else None,
                    'right': tcBorders.find(qn('w:right')) if tcBorders is not None else None,
                }

                for side, border in border_sides.items():
                    if border is not None:
                        color = border.get(qn('w:color'), "").upper()  # Ensure uppercase and no #
                        size = border.get(qn('w:sz'))
                        style = border.get(qn('w:val'))
                    else:
                        color, size, style = "none", "none", "none"

                    # Convert size from eighths of a point to points
                    size_in_pt = str(round(float(size) / 8, 3)) if size and size != "none" else "none"

                    # Get expected properties for the current cell and side
                    expected_properties = expected_sides[cell_idx][side]

                    # Assertions
                    assert color == expected_properties["color"], (
                        f"Color mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['color']}, got {color}"
                    )
                    assert size_in_pt == expected_properties["size"], (
                        f"Size mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['size']}, got {size_in_pt}"
                    )
                    assert style == expected_properties["style"], (
                        f"Style mismatch for {side} in row {row_idx} column {column_idx}: "
                        f"expected {expected_properties['style']}, got {style}"
                    )

                cell_idx += 1

    def test_unbalanced_table(self):
        # A table with more td elements in latter rows than in the first
        self.document.add_heading('Test: Handling unbalanced tables', level=1)

        html_unbalanced_table = """
            <table>
            <tr><td>Hello</td></tr>
            <tr><td>One</td><td>Two</td></tr>
            </table>
        """
        self.parser.add_html_to_document(html_unbalanced_table, self.document)
        document = self.parser.parse_html_string(html_unbalanced_table)

        # Get the last table added to the document
        tables = document.tables
        assert len(tables) == 1

        # Docx will autofit all cells
        table = tables[0]
        assert len(table.rows) == 2
        assert len(table.rows[0].cells) == 2
        assert len(table.rows[1].cells) == 2

        assert table.rows[0].cells[0].text.strip() == "Hello"
        assert table.rows[1].cells[0].text.strip() == "One"
        assert table.rows[1].cells[1].text.strip() == "Two"

    def test_html_comment_rendering(self):
        self.document.add_heading("Test: HTML Comment Rendering", level=1)

        html_with_comment = """
        <p>Hello</p>
        <!-- This is a comment -->
        <p>World</p>
        """

        # Process document using parser
        self.parser.options['html-comments'] = True
        self.parser.add_html_to_document(html_with_comment, self.document)
        document = self.parser.parse_html_string(html_with_comment)

        # Extract all paragraph texts
        paragraph_texts = [p.text.strip() for p in document.paragraphs]

        # Expected comment result
        expected_comment = "# This is a comment"

        # Assert the comment paragraph exists
        assert any(
            expected_comment == text for text in paragraph_texts
        ), f"Expected comment '{expected_comment}' to appear in the document, but it was not found."

        # (Optional) Check styling if needed: green color or italic
        comment_paragraph = next(
            (p for p in document.paragraphs if p.text.strip() == expected_comment),
            None
        )
        assert comment_paragraph is not None, "Comment paragraph not found for style checks."

        comment_run = comment_paragraph.runs[0]

        # italic assertion
        assert comment_run.italic, "HTML comment should be italic."

        # color assertion (dark-ish green #008000)
        expected_rgb = parse_color("#008000")
        assert (
            comment_run.font.color.rgb is not None
            and comment_run.font.color.rgb == RGBColor(*expected_rgb)
        ), f"Comment run color should be green ({expected_rgb})."

    def test_emojis_and_special_characters(self):
        emojis_and_special_chars_html_example = """
        <html>
            <body>
                <p>Emoji Test: 😊🔥🎉</p>
                <p>HTML Entities: &amp; &lt; &gt; &copy; &reg;</p>
                <p>Math Symbols: ∑ π √</p>
            </body>
        </html>
        """

        self.document.add_heading(
            'Test: Emojis and Special Characters',
            level=1
        )
        # Add on document for human validation
        self.parser.add_html_to_document(emojis_and_special_chars_html_example, self.document)

        document = self.parser.parse_html_string(emojis_and_special_chars_html_example)
        doc_text = " ".join([p.text for p in document.paragraphs])

        # Check if all expected elements exist in the DOCX
        assert "😊" in doc_text, "Emoji '😊' is missing"
        assert "🔥" in doc_text, "Emoji '🔥' is missing"
        assert "🎉" in doc_text, "Emoji '🎉' is missing"
        assert "&" in doc_text, "HTML entity '&amp;' did not convert correctly"
        assert "<" in doc_text, "HTML entity '&lt;' did not convert correctly"
        assert ">" in doc_text, "HTML entity '&gt;' did not convert correctly"
        assert "©" in doc_text, "HTML entity '&copy;' did not convert correctly"
        assert "®" in doc_text, "HTML entity '&reg;' did not convert correctly"
        assert "∑" in doc_text, "Math symbol '∑' is missing"
        assert "π" in doc_text, "Math symbol 'π' is missing"
        assert "√" in doc_text, "Math symbol '√' is missing"

    def test_ordered_list(self):
        self.document.add_heading("Test: Ordered List", level=1)

        ordered_list_html_example = """
        <ol>
          <li>first list, item 1</li>
          <p>Paragraph inserted between items</p>
          <li>first list, item 2</li>
          <li><p>first list, item 3 within a paragraph</p></li>
        </ol>
        """
        self.parser.add_html_to_document(ordered_list_html_example, self.document)
        document = self.parser.parse_html_string(ordered_list_html_example)

        # Extract paragraphs with 'ListNumber' style (ordered list)
        ordered_list_paragraphs = [
            p for p in document.paragraphs if p.style.name == "List Number"
        ]

        # Expected items in order
        expected_items = [
            "first list, item 1",
            "first list, item 2",
            "first list, item 3 within a paragraph",
        ]

        assert len(ordered_list_paragraphs) >= len(
            expected_items
        ), f"Expected at least {len(expected_items)} ordered list items, found {len(ordered_list_paragraphs)}"

        for i, expected_text in enumerate(expected_items):
            actual_text = ordered_list_paragraphs[i].text.strip()
            assert (
                actual_text == expected_text
            ), f"Expected ordered list item '{expected_text}', but got '{actual_text}'"

    def test_unordered_list(self):
        self.document.add_heading("Test: Unordered List", level=1)

        unordered_list_html_example = """
        <ul>
          <li>Unorderd list</li>
          <p>Paragraph inserted between items</p>
          <li>with circle markers</li>
          <li><p>last option</p></li>
        </ul>
        """
        self.parser.add_html_to_document(unordered_list_html_example, self.document)
        document = self.parser.parse_html_string(unordered_list_html_example)

        # Extract paragraphs with 'ListBullet' style (unordered list)
        unordered_list_paragraphs = [
            p for p in document.paragraphs if p.style.name == "List Bullet"
        ]

        # Expected unordered items
        expected_items = ["Unorderd list", "with circle markers", "last option"]

        assert len(unordered_list_paragraphs) >= len(
            expected_items
        ), f"Expected at least {len(expected_items)} unordered list items, found {len(unordered_list_paragraphs)}"

        for expected_text in expected_items:
            assert any(
                expected_text == p.text.strip() for p in unordered_list_paragraphs
            ), f"Unordered list item '{expected_text}' not found in List Bullet paragraphs"

    def test_table_rowspan_and_colspan(self):
        self.document.add_heading("Test: Table rowspan and colspan", level=1)

        rowspan_and_colspan_html_example = """
        <table border="1" cellspacing="0" cellpadding="4">
          <tr>
            <th rowspan="2" style="vertical-align:middle;text-align:center">Region</th>
            <th colspan="2" style="vertical-align:middle;text-align:center">Sales (CHF millions)</th>
          </tr>
          <tr>
            <th>2024</th>
            <th>2023</th>
          </tr>
          <tr>
            <td>United States</td>
            <td>22,456</td>
            <td>20,892</td>
          </tr>
          <tr>
            <td>Europe</td>
            <td>8,147</td>
            <td>7,634</td>
          </tr>
        </table>"""
        self.parser.table_style = 'Table Grid'
        self.parser.add_html_to_document(rowspan_and_colspan_html_example, self.document)
        document = self.parser.parse_html_string(rowspan_and_colspan_html_example)

        # Find the first table
        table = document.tables[0]

        # Assertions on structure
        assert len(table.rows) == 4, "Table should have 4 rows"
        assert len(table.columns) == 3, "Table should have 3 columns"

        # Cell (0, 0): Region (rowspan=2)
        assert "Region" in table.cell(0, 0).text
        assert table.cell(0, 0)._tc.tcPr.vMerge is not None, "Table Cell (0, 0) is not vertically merged"
        assert table.cell(1, 0)._tc.tcPr.vMerge is not None, "Table Cell (1, 0) is not vertically merged"

        # Cell (0, 1): Sales (CHF millions) (colspan=2)
        sales_cell = table.cell(0, 1)
        assert "Sales" in sales_cell.text
        assert sales_cell._tc.tcPr.gridSpan is not None, "Table Cell (0, 1) is not horizontally merged"
        assert int(sales_cell._tc.tcPr.gridSpan.val) == 2, "Table Cell (0, 2) is not horizontally merged"

        # Verify unmerged data cells
        assert "2024" in table.cell(1, 1).text
        assert "2023" in table.cell(1, 2).text
        assert "United States" in table.cell(2, 0).text
        assert "22,456" in table.cell(2, 1).text
        assert "20,892" in table.cell(2, 2).text
        assert "Europe" in table.cell(3, 0).text
        assert "8,147" in table.cell(3, 1).text
        assert "7,634" in table.cell(3, 2).text

    def test_complex_colspan_rowspan_combinations(self):
        self.document.add_heading('Test: Complex Colspan and Rowspan Combinations', level=1)

        complex_table_html = """
        <table border="1">
            <tr>
                <td rowspan="2">A1-A2</td>
                <td colspan="3">B1-D1</td>
                <td>E1</td>
            </tr>
            <tr>
                <td>B2</td>
                <td colspan="2" rowspan="2">C2-D3</td>
                <td rowspan="3">E2-E4</td>
            </tr>
            <tr>
                <td colspan="2">A3-B3</td>
            </tr>
            <tr>
                <td>A4</td>
                <td>B4</td>
                <td>C4</td>
                <td>D4</td>
            </tr>
        </table>
        """

        try:
            self.parser.table_style = 'Table Grid'
            self.parser.add_html_to_document(complex_table_html, self.document)
            document = self.parser.parse_html_string(complex_table_html)

            tables = document.tables
            assert len(tables) == 1, "Should create a table"

            table = tables[0]
            assert len(table.rows) == 4, f"Expected 4 rows, but got {len(table.rows)} rows"
            assert len(table.columns) == 5, f"Expected 5 columns, but got {len(table.columns)} columns"

            assert "A1-A2" in table.cell(0, 0).text, "First merged cell content is incorrect"
            assert "B1-D1" in table.cell(0, 1).text, "Second merged cell content is incorrect"

        except IndexError as e:
            self.fail(f"Complex table processing failed with IndexError: {e}")
        except Exception as e:
            self.fail(f"Processing complex table failed with unexpected error: {e}")

    def test_extreme_colspan_rowspan_cases(self):
        """ Test extreme colspan and rowspan cases """
        self.document.add_heading('Test: Extreme Colspan and Rowspan Cases', level=1)

        extreme_table_html = """
        <table border="1">
            <tr>
                <td colspan="10">Extreme colspan cell</td>
            </tr>
            <tr>
                <td rowspan="5">Extreme rowspan cell</td>
                <td colspan="9">Extreme colspan cell</td>
            </tr>
            <tr>
                <td colspan="3">Col 1-3</td>
                <td colspan="3">Col 4-6</td>
                <td colspan="3">Col 7-9</td>
            </tr>
            <tr>
                <td>1</td><td>2</td><td>3</td><td>4</td><td>5</td><td>6</td><td>7</td><td>8</td><td>9</td>
            </tr>
            <tr>
                <td colspan="2">1-2</td><td colspan="2">3-4</td><td colspan="2">5-6</td><td colspan="3">7-9</td>
            </tr>
            <tr>
                <td colspan="9">The last row should span all columns</td>
            </tr>
        </table>
        """

        try:
            self.parser.table_style = 'Table Grid'
            self.parser.add_html_to_document(extreme_table_html, self.document)
            document = self.parser.parse_html_string(extreme_table_html)

            tables = document.tables
            assert len(tables) == 1, "Should create a table"

            table = tables[0]

            assert len(table.rows) == 6, f"Expected 6 rows, but got {len(table.rows)} rows"
            assert len(table.columns) == 10, f"Expected 10 columns, but got {len(table.columns)} columns"

            assert "Extreme colspan cell" in table.cell(0, 0).text, "First cell content is incorrect"
            assert "Extreme rowspan cell" in table.cell(1, 0).text, "Second cell content is incorrect"

        except IndexError as e:
            self.fail(f"Extreme table processing failed with IndexError: {e}")
        except Exception as e:
            self.fail(f"Processing extreme table failed with unexpected error: {e}")

    def test_nested_styles_on_multiple_tags(self):
        """ Test nested styles on multiple tags """
        self.document.add_heading('Test: Test nested styles on multiple tags', level=1)

        nested_styles_html = """
        <h3 style="color: red; font-size:24px">Title Text</h3>
        <div style="background-color: black; color: #fff; font-size:1rem; text-align: center">
            Div Text
            <p style="color: lightgreen; text-align: center">
                P Text
            </p>

            <br>

            <ol>
                <li style="color: lightblue; font-size: 12px">Li Text 1</li>
                <li style="color: lightyellow; font-size: 8px">Li Text 2</li>
            <ol>
        </div>
        """

        self.parser.add_html_to_document(nested_styles_html, self.document)
        document = self.parser.parse_html_string(nested_styles_html)

        # -------- H3 ----------
        h3_paragraphs = [p for p in document.paragraphs if 'Title Text' in p.text]
        assert len(h3_paragraphs) == 1
        h3_run = h3_paragraphs[0].runs[0]
        assert h3_run.text == 'Title Text'
        assert h3_run.font.color.rgb == Color['red'].value
        assert h3_run.font.size is not None

        # -------- Div ----------
        div_paragraphs = [p for p in document.paragraphs if 'Div Text' in p.text]
        assert len(div_paragraphs) == 1
        div_run = div_paragraphs[0].runs[0]
        assert div_run.text.strip() == 'Div Text'
        assert div_run.font.color.rgb == Color['white'].value
        assert div_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

        # -------- P inside div ----------
        p_paragraphs = [p for p in document.paragraphs if 'P Text' in p.text]
        assert len(p_paragraphs) == 1
        p_run = p_paragraphs[0].runs[0]
        assert p_run.text.strip() == 'P Text'
        assert p_run.font.color.rgb == Color['lightgreen'].value
        assert p_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

        # -------- List items ----------
        li1_paragraphs = [p for p in document.paragraphs if 'Li Text 1' in p.text]
        assert len(li1_paragraphs) == 1
        li1_run = li1_paragraphs[0].runs[0]
        assert li1_run.text.strip() == 'Li Text 1'
        assert li1_run.font.color.rgb == Color['lightblue'].value
        assert li1_run.font.size is not None

        li2_paragraphs = [p for p in document.paragraphs if 'Li Text 2' in p.text]
        assert len(li2_paragraphs) == 1
        li2_run = li2_paragraphs[0].runs[0]
        assert li2_run.text.strip() == 'Li Text 2'
        assert li2_run.font.color.rgb == Color['lightyellow'].value
        assert li2_run.font.size is not None

    def test_basic_class_mapping(self):
        """Test that CSS classes are mapped to Word styles"""
        self.document.add_heading("Test: Test Basic Class Mapping", level=1)
        style_map = {
            "custom-style": "Quote",
        }

        html = '<p class="custom-style">Test paragraph</p>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.options["style-map"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # Verify paragraph uses the mapped style
        self.assertEqual(doc.paragraphs[0].style.name, "Quote")

    def test_multiple_classes(self):
        """Test that first matching class in style_map wins"""
        self.document.add_heading(
            "Test: Test that first matching class in style_map wins", level=1
        )
        style_map = {
            "first": "Heading 2",
            "second": "Heading 3",
        }

        html = '<p class="second first">Test</p>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.options["style-map"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # Should use first matching class found
        self.assertIn(doc.paragraphs[0].style.name, ["Heading 2", "Heading 3"])

    def test_unmapped_class_uses_default(self):
        """Test that unmapped classes fall back to default behavior"""
        self.document.add_heading(
            "Test: Test that unmapped classes fall back to default behavior", level=1
        )
        style_map = {
            "mapped": "Heading 450",
        }

        html = '<p class="unmapped">Test</p>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map, default_paragraph_style=None)
        parser.options["style-map"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # Should use default Word 'Normal' style
        self.assertEqual(doc.paragraphs[0].style.name, "Normal")

    def test_h1_override(self):
        """Test overriding default h1 style"""
        self.document.add_heading("Test: Test H1 Override", level=1)
        tag_overrides = {
            "h1": "Heading 2",
        }

        html = "<h1>Test Heading</h1>"

        doc = Document()
        parser = HtmlToDocx(tag_style_overrides=tag_overrides)
        parser.options["tag-override"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # h1 should use Heading 2 instead of default Heading 1
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 2")

    def test_class_overrides_tag_override(self):
        """Test that class mapping has priority over tag override"""
        self.document.add_heading(
            "Test: Test class mapping priority over tag override", level=1
        )
        style_map = {"custom": "Heading 3"}
        tag_overrides = {"h1": "Heading 2"}

        html = '<h1 class="custom">Test</h1>'

        doc = Document()
        parser = HtmlToDocx(style_map=style_map, tag_style_overrides=tag_overrides)
        parser.options["style-map"] = True
        parser.options["tag-override"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # Class should win over tag override
        self.assertEqual(doc.paragraphs[0].style.name, "Heading 3")

    def test_normal_default(self):
        """Test that Normal is used as default by default"""
        self.document.add_heading(
            "Test: Test that Normal style is used as default", level=1
        )
        html = "<p>Test paragraph</p>"

        doc = Document()
        parser = HtmlToDocx()  # default_paragraph_style=None by default
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[0].style.name, "Normal")

    def test_custom_default(self):
        """Test setting custom default paragraph style"""
        self.document.add_heading("Test: Test custom default paragraph style", level=1)
        html = "<p>Test paragraph</p>"

        doc = Document()
        parser = HtmlToDocx(default_paragraph_style="Heading 1")
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[0].style.name, "Heading 1")

    def test_none_default_uses_normal(self):
        """Test that None uses Word's default Normal style"""
        self.document.add_heading(
            "Test: Test default of None will use 'Normal' as default style", level=1
        )
        html = "<p>Test paragraph</p>"

        doc = Document()
        parser = HtmlToDocx(default_paragraph_style=None)
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[0].style.name, "Normal")

    def test_fontweight_bold(self):
        """Test font-weight bold"""
        html = '<p><span style="font-weight: bold">Bold text</span></p>'
        self.document.add_heading("Test: Test Font-Weight bold", level=1)

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertTrue(run.font.bold)

    def test_fontstyle_italic(self):
        """Test font-style italic"""
        html = '<p><span style="font-style: italic">Italic text</span></p>'
        self.document.add_heading("Test: Test Font-Style italics", level=1)

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertTrue(run.font.italic)

    # def test_textdecoration(self):
    #     """Test text-decoration"""
    #     # 16px = 12pt
    #     html = '<p><span style="text-decoration: underline wavy blue 16px">An underlined, blue wavy text.</span></p>'
    #     self.document.add_heading("Test: Test Text-Decoration", level=1)

    #     doc = Document()
    #     parser = HtmlToDocx()
    #     parser.add_html_to_document(html, self.document)
    #     parser.add_html_to_document(html, doc)

    #     run = doc.paragraphs[0].runs[0]
    #     blue_font = run.font.color.rgb == Color["blue"].value
    #     is_underlined = True if run.font.underline is not None else False
    #     is_underline_wavy = True if run.font.underline == WD_UNDERLINE.WAVY else False
    #     result_list = [blue_font, is_underlined, is_underline_wavy]
    #     self.assertTrue(all(result_list))

    def test_fontweight_none(self):
        """Test None as font-weight Value"""
        html = '<p><span style="font-weight: None">Regular text</span></p>'
        self.document.add_heading("Test: Test font-weight as None", level=1)

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertTrue(run.font.bold is not True)

    def test_fontstyle_none(self):
        """Test font-style italic"""
        html = '<p><span style="font-style: none">Italic text</span></p>'
        self.document.add_heading("Test: Test font-style None", level=1)

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertTrue(run.font.italic is not True)

    # def test_textdecoration_none(self):
    #     """Test text-decoration as None"""
    #     # 16px = 12pt
    #     html = '<p><span style="text-decoration: none;">An regular boring text with no decorations...</span></p>'
    #     self.document.add_heading("Test: Test Text-Decoration None", level=1)

    #     doc = Document()
    #     parser = HtmlToDocx()
    #     parser.add_html_to_document(html, self.document)
    #     parser.add_html_to_document(html, doc)

    #     run = doc.paragraphs[0].runs[0]
    #     black_font = run.font.color.rgb == Color["black"].value
    #     is_not_underlined = True if run.font.underline is None else True
    #     is_not_underline_wavy = True if run.font.underline is None else False
    #     results = [black_font, is_not_underlined, is_not_underline_wavy]
    #     print(results)
    #     self.assertTrue(all(results))

    def test_paragraph_inline_styles(self):
        """Test inline styles on paragraph elements"""
        html = '<p style="color: blue; font-size: 14pt">Blue 14pt paragraph</p>'
        self.document.add_heading("Test: Test paragraph inline styles", level=1)

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertIsNotNone(run.font.color.rgb)
        self.assertEqual(run.font.size, Pt(14))

    def test_important_overrides_normal(self):
        """Test that !important styles override normal styles"""
        self.document.add_heading("Test: Test !important override", level=1)
        html = """
        <p>
            <span style="color: gray">
                Gray text with <span style="color: red !important">red important</span>.
            </span>
        </p>
        """

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # The "red important" run should have red color
        # (exact run index may vary based on whitespace handling)
        run = doc.paragraphs[1].runs[2]
        self.assertEqual(run.font.color.rgb, Color["red"].value)

    def test_important_conflict_last_wins(self):
        """Test conflict when both styles have !important"""
        self.document.add_heading("Test: Test Last !important override", level=1)
        html = """
        <p>
            <span style="color: BLUE !important">
                Blue text with <span style="color: red !important">red important</span>.
            </span>
        </p>
        """
        # html = '<p><span style="color: BLUE !important">Blue text with <span style="color: red !important">red important</span>.</span></p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # The "red important" run should have red color
        # (exact run index may vary based on whitespace handling)

        # You can see the index by uncommented the following block.  The multi-line html currently in use creates 2
        # paragraphs.  The first has 1 empty run.  The 2nd has 6 runs (3 of which are empty due to whitespace).
        # The single line version (commented out above), creates a single paragraph with 3 runs.

        # print("Paragraph count:", len(doc.paragraphs))
        # for i, para in enumerate(doc.paragraphs):
        #     print(f"\nParagraph {i}:")
        #     print("  Run count:", len(para.runs))

        #     for j, temp in enumerate(para.runs):
        #         print(f"    Run {j}: '{temp.text}' Color: {temp.font.color.rgb}")

        run = doc.paragraphs[1].runs[2]

        self.assertEqual(run.font.color.rgb, Color["red"].value)

    def test_important_on_paragraph(self):
        """Test !important on paragraph inline style"""
        self.document.add_heading(
            "Test: Test !important override for paragraph", level=1
        )
        html = '<p style="color: blue !important">Blue important</p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertIsNotNone(run.font.color.rgb)

    def test_multi_paragraph_code_block(self):
        """Test that all paragraphs in code block maintain style"""
        self.document.add_heading("Test: Test multi-paragraph code block", level=1)
        style_map = {
            "code-block": "No Spacing",  # Using built-in style
        }

        html = """
        <div class="code-block">
            <p>First line of code</p>
            <p>Second line of code</p>
            <p>Third line of code</p>
        </div>
        """

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.options["style-map"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # All three paragraphs should have the code-block style
        self.assertEqual(doc.paragraphs[1].style.name, "No Spacing")
        self.assertEqual(doc.paragraphs[2].style.name, "No Spacing")
        self.assertEqual(doc.paragraphs[3].style.name, "No Spacing")

    def test_numbered_headings(self):
        """Test numbered heading classes"""
        self.document.add_heading("Test: Test Numbered heading (sorta)", level=1)
        style_map = {
            "numbered-heading-1": "Heading 3",
            "numbered-heading-2": "Heading 4",
            "numbered-heading-3": "Heading 5",
        }

        html = """
        <h1 class="numbered-heading-1" style="color: red">1.0 Introduction</h1>
        <h2 class="numbered-heading-2" style="color: red">1.1 Overview</h2>
        <h3 class="numbered-heading-3" style="color: red">1.1.1 Details</h3>
        """

        doc = Document()
        parser = HtmlToDocx(style_map=style_map)
        parser.options["style-map"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        self.assertEqual(doc.paragraphs[1].style.name, "Heading 3")
        self.assertEqual(doc.paragraphs[2].style.name, "Heading 4")
        self.assertEqual(doc.paragraphs[3].style.name, "Heading 5")

    def test_basic_html_still_works(self):
        """Test that basic HTML conversion works without new features"""
        self.document.add_heading(
            "Test: Test Basic HTML still works after changes", level=1
        )
        html = "<p>Simple paragraph</p><h3> and here we have heading 3</h3>"

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        self.assertEqual(len(doc.paragraphs), 2)
        self.assertEqual(doc.paragraphs[1].style.name, "Heading 3")

    def test_existing_span_styles_work(self):
        """Test that existing <span style="..."> still works"""
        self.document.add_heading("Test: Test Existing span styles", level=1)
        html = '<p><span style="color: #FF0000">Red text</span></p>'

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        run = doc.paragraphs[0].runs[0]
        self.assertIsNotNone(run.font.color.rgb)

    def test_bold_italic_tags_work(self):
        """Test that <b>, <i>, <u> tags still work"""
        self.document.add_heading(
            "Test: bold, itatlic, and underline tags to ensure they still work", level=1
        )
        html = "<p><b>Bold</b> <i>Italic</i> <u>Underline</u></p>"

        doc = Document()
        parser = HtmlToDocx()
        parser.add_html_to_document(html, doc)
        parser.add_html_to_document(html, self.document)

        # Using in memory "doc" for assertion to isolate test
        # Find runs with the specific formatting (spaces create extra runs, so we can't rely on indices)
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.font.bold]
        italic_runs = [r for r in runs if r.font.italic]
        underline_runs = [r for r in runs if r.font.underline]

        self.assertTrue(len(bold_runs) > 0, "Should have at least one bold run")
        self.assertTrue(len(italic_runs) > 0, "Should have at least one italic run")
        self.assertTrue(
            len(underline_runs) > 0, "Should have at least one underline run"
        )

    def test_nonexistent_style_graceful_failure(self):
        """Test that non-existent styles don't crash"""
        self.document.add_heading(
            "Test: Test crash protection when style doesn't exist", level=1
        )
        style_map = {
            "custom": "NonExistentStyle",
        }

        html = '<p class="custom">Test</p>'

        parser = HtmlToDocx(style_map=style_map)
        parser.options["style-map"] = True

        # Should not raise exception
        try:
            parser.add_html_to_document(html, self.document)
            success = True
        except Exception:
            success = False

        self.assertTrue(success)

    def test_empty_style_map(self):
        """Test with empty style_map"""
        self.document.add_heading("Test: Test empty style map", level=1)
        html = '<p class="anything">Test</p>'

        doc = Document()
        parser = HtmlToDocx(style_map={})
        parser.options["style-map"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        # Should use default (Normal)
        self.assertEqual(doc.paragraphs[0].style.name, "Normal")

    def test_none_style_map(self):
        """Test with None style_map"""
        self.document.add_heading("Test: Test None as style map", level=1)
        html = "<p>Test</p>"

        doc = Document()
        parser = HtmlToDocx(style_map=None)
        parser.options["style-map"] = True
        parser.add_html_to_document(html, self.document)
        parser.add_html_to_document(html, doc)

        self.assertEqual(len(doc.paragraphs), 1)


if __name__ == "__main__":
    unittest.main()
