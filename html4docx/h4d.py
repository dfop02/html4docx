import argparse
import logging
import os
import re
from io import BytesIO
from html.parser import HTMLParser

import docx
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from functools import lru_cache

from html4docx import constants
from html4docx import utils
from html4docx.metadata import Metadata

class HtmlToDocx(HTMLParser):
    """
        Class to convert HTML to Docx
        source: https://docs.python.org/3/library/html.parser.html
    """
    def __init__(self, style_map=None, tag_style_overrides=None, default_paragraph_style="Normal"):
        super().__init__()
        self.options = dict(constants.DEFAULT_OPTIONS)
        self.table_row_selectors = constants.DEFAULT_TABLE_ROW_SELECTORS
        self.table_style = constants.DEFAULT_TABLE_STYLE
        self.paragraph_span_styles = {}  # paragraph_id -> set(run_indices)
        self.style_map = style_map or constants.DEFAULT_STYLE_MAP
        self.tag_style_overrides = tag_style_overrides or constants.DEFAULT_TAG_OVERRIDES
        self.default_paragraph_style = default_paragraph_style or constants.DEFAULT_PARAGRAPH_STYLE

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        self.doc = document if document else Document()
        self.bs = self.options['fix-html'] # whether or not to clean with BeautifulSoup
        self.paragraph = None
        self.run = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0
        self.bookmark_id = 0
        self.in_li = False
        self.list_restart_counter = 0
        self.current_ol_num_id = None
        self._list_num_ids = {}

        # NEW: Set style map & tag overrides according to options

        self.use_styles = False if self.options["styles"] is False else self.options["style-map"]
        self.use_tag_overrides = self.options["tag-override"]
        # NEW: Style tracking variables
        self.pending_div_style = None
        self.pending_character_style = None
        self.pending_inline_styles = None
        self.pending_important_styles = None

    @property
    def metadata(self) -> dict[str, any]:
        if not hasattr(self, '_metadata'):
            self._metadata = Metadata(self.doc)
        return self._metadata

    @property
    def include_tables(self) -> bool:
        return self.options.get('tables', True)

    @property
    def include_images(self) -> bool:
        return self.options.get('images', True)

    @property
    def include_styles(self) -> bool:
        return self.options.get('styles', True)

    @property
    def include_html_comments(self) -> bool:
        return self.options.get('html-comments', False)

    @property
    def include_stylemap(self) -> bool:
        return self.options.get("style-map", True)

    @property
    def include_tagoverrides(self) -> bool:
        return self.options.get("tag-override", True)

    def save(self, destination) -> None:
        """Save the document to a file path or BytesIO object."""
        if isinstance(destination, str):
            destination, _ = os.path.splitext(destination)
            self.doc.save(f'{destination}.docx')
        elif isinstance(destination, BytesIO):
            self.doc.save(destination)
        else:
            raise TypeError('destination must be a str path or BytesIO object')

    def copy_settings_from(self, other):
        """Copy settings from another instance of HtmlToDocx"""
        self.table_style = other.table_style

        # NEW: Copy extended settings if present
        if hasattr(other, "style_map"):
            self.style_map = other.style_map
            self.tag_style_overrides = other.tag_style_overrides
            self.default_paragraph_style = other.default_paragraph_style

    def get_word_style_for_element(self, tag, attrs):
        """
            Determine the Word style to use for an HTML element.

            Priority order:
            1. CSS class from style_map (if present)
            2. Tag override from tag_style_overrides
            3. Default behavior

            Args:
            tag: HTML tag name (e.g., 'h1', 'p', 'div')
            attrs: Dictionary of HTML attributes

        Returns:
            str or None: Word style name to apply, or None for default behavior
        """
        # Priority 1: Check if element has a class attribute mapped in style_map
        if "class" in attrs:
            # html class can be multiple classes separated by space
            classes = attrs["class"].split()
            for cls in classes:
                if cls in self.style_map:
                    return self.style_map[cls]
        # Priority 2, tag override
        if tag in self.tag_style_overrides:
            return self.tag_style_overrides[tag]

        # Priority 3, default behavior.
        return None

    def apply_style_to_paragraph(self, paragraph, style_name):
        """
        Apply a Word style to a paragraph by style name.

        Args:
            paragraph: python-docx Paragraph object
            style_name (str): Name of the Word style to apply

        Returns:
            bool: True if style was applied successfully, False otherwise
        """
        try:
            paragraph.style = style_name
            return True
        except KeyError:
            # Style doesn't exist in document
            print(
                f"Warning: Style '{style_name}' not found in document. Using default."
            )
            return False

    def apply_style_to_run(self, style_name):
        """
        Apply a Word character style to a run by style name.

        Args:
            run: python-docx Run object
            style_name (str): Name of the Word character style to apply

        Returns:
            bool: True if style was applied successfully, False otherwise
        """
        try:
            self.run.style = style_name
            return True
        except KeyError:
            print(f"Warning: Character style '{style_name}' not found in document.")
            return False
        except ValueError as e:
            if "need type CHARACTER" in str(e):
                print(
                    f"Warning: '{style_name}' is a paragraph style, not a character style."
                )
                print(
                    "For inline elements like <code>, please create a character style in Word."
                )
            return False

    def parse_inline_styles(self, style_string):
        """
        Parse inline CSS styles and separate normal styles from !important ones.

        Args:
            style_string (str): CSS style string (e.g., "color: red; font-size: 12px !important")

        Returns:
            tuple: (normal_styles dict, important_styles dict)
        """
        normal_styles = {}
        important_styles = {}

        if not style_string:
            return normal_styles, important_styles

        # Parse style string into individual declarations
        style_dict = utils.parse_dict_string(style_string)

        for prop, value in style_dict.items():
            # Check if value has !important flag
            if "!important" in value.lower():
                # Remove !important flag and store in important_styles
                clean_value = utils.remove_important_from_style(value)
                important_styles[prop] = clean_value
            else:
                normal_styles[prop] = value

        return normal_styles, important_styles

    def apply_inline_styles_to_run(self, styles_dict):
        """
        Apply inline CSS styles to a run.

        Supports: color, background-color, font-size, font-weight, font-style,
                text-decoration, font-family

        Args:
            run: python-docx Run object
            styles_dict: Dictionary of CSS properties and values
        """
        if not styles_dict:
            return

        # Apply color
        if "color" in styles_dict:
            try:
                colors = utils.parse_color(styles_dict["color"])
                self.run.font.color.rgb = RGBColor(*colors)
            except Exception:
                pass

        # Apply font-size
        if "font-size" in styles_dict:
            try:
                font_size = utils.adapt_font_size(styles_dict["font-size"])
                if font_size:
                    self.run.font.size = utils.unit_converter(font_size)
            except Exception:
                pass

        # Apply font-weight (bold)
        if "font-weight" in styles_dict:
            weight = styles_dict["font-weight"].lower()
            if weight in ["bold", "bolder", "700", "800", "900"]:
                self.run.font.bold = True
            elif weight in ["normal", "400"]:
                self.run.font.bold = False

        # Apply font-style (italic)
        if "font-style" in styles_dict:
            style = styles_dict["font-style"].lower()
            if style == "italic" or style == "oblique":
                self.run.font.italic = True
            elif style == "normal":
                self.run.font.italic = False

        # Apply font-family
        if "font-family" in styles_dict:
            font_family = (
                styles_dict["font-family"].split(",")[0].strip().strip('"').strip("'")
            )
            self.run.font.name = font_family

    def get_cell_html(self, soup):
        """
        Returns string of td element with opening and closing <td> tags removed
        Cannot use find_all as it only finds element tags and does not find text which
        is not inside an element
        """
        return ' '.join([str(i) for i in soup.contents])

    def set_cell_background(self, cell, color):
        """Set the background color of a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color.lstrip('#'))
        tcPr.append(shd)

    def set_cell_borders(self, cell, styles):
        """
        Apply custom borders to a table cell using CSS-like styles.

        Args:
            cell: The table cell to style.
            styles (dict): A dictionary of CSS-like border styles.
                          Example: {"border": "1px 5px 10px 20px", "border-style": "solid", "border-color": "#FF0000"}
        sources:
            Border attributes: http://officeopenxml.com/WPtableBorders.php
        """
        tc = cell._tc  # Get the table cell element
        tcPr = tc.get_or_add_tcPr()  # Get or add the table cell properties

        # Default border properties
        default_size = constants.DEFAULT_BORDER_SIZE
        default_color = constants.DEFAULT_BORDER_COLOR
        default_style = constants.DEFAULT_BORDER_STYLE
        borders = constants.default_borders()
        border_styles = constants.BORDER_STYLES
        keywords = constants.BORDER_KEYWORDS
        border_sides = ("top", "right", "bottom", "left")
        border_width_pattern = re.compile(r'^[0-9]*\.?[0-9]+(px|pt|cm|in|rem|em|%)$')

        def parse_border_style(value: str) -> str:
            """Parses border styles to match word standart"""
            return constants.BORDER_STYLES[value] if value in constants.BORDER_STYLES.keys() else 'none'

        def check_unit_keywords(value: str) -> str:
            """Convert medium, thin, thick keywords to numeric values (px)"""
            lower_val = value.lower()
            return keywords.get(lower_val, value)

        @lru_cache(maxsize=None)
        def border_unit_converter(unit_value: str):
            """Convert multiple units to pt that is used on Word table cell border"""
            unit_value = utils.remove_important_from_style(unit_value)
            unit_value = check_unit_keywords(unit_value)

            # Return default if no value or empty
            if not unit_value or unit_value == '':
                return default_size

            unit = re.sub(r'[0-9\.]+', '', unit_value)
            value = float(re.sub(r'[a-zA-Z\!\%]+', '', unit_value))  # Allow float values

            if unit == 'px':
                result = value * 0.75  # 1 px = 0.75 pt
            elif unit == 'cm':
                result = value * 28.35  # 1 cm = 28.35 pt
            elif unit == 'in':
                result = value * 72  # 1 inch = 72 pt
            elif unit == 'pt':
                result = value # default is pt
            elif unit == 'rem' or unit == 'em':
                result = value * 12  # Assuming 1rem/em = 16px, converted to pt
            elif unit == '%':
                result = constants.MAX_INDENT * (value / 100)
            else:
                return None  # Unsupported units return None

            return result

        def parse_border_value(value: str):
            """
            Parses a border value like:
            '1px solid #000000', 'solid 1px red', or '#000000 medium dashed' in any order.
            """
            parts = value.split()

            # Return all default if there is only 'none' or empty
            if (len(parts) == 1 and parts[0] == 'none') or (not value or value.strip() == ''):
                return default_size, default_style, default_color

            size = None
            style = default_style
            color = default_color

            for part in parts:
                clean_part = utils.remove_important_from_style(part).lower()

                # Detect size (units or keywords)
                if border_width_pattern.match(clean_part) or clean_part in keywords:
                    size = border_unit_converter(clean_part) or default_size
                    continue

                # Detect style
                if clean_part in border_styles:
                    style = parse_border_style(clean_part)
                    continue

                # Detect color
                if utils.is_color(clean_part):
                    color = utils.parse_color(clean_part, return_hex=True)
                    continue

            # If only style or color was given without size, we need to apply 1pt size to render it
            if len(parts) >= 1 and size is None:
                size = 1.0

            return size, style, color

        for css_prop, css_value in styles.items():
            if not css_prop.startswith("border"):
                continue

            # Case 1: 'border' shorthand applies to all sides
            if css_prop == "border":
                values = css_value.split()
                if len(values) in (1, 3):
                    # Single value or full triple — parse all parts in any order
                    size, style, color = parse_border_value(css_value)
                    for side in border_sides:
                        borders[side].update({"size": size, "style": style, "color": color})
                elif len(values) == 2:
                    # Two widths (top/bottom, left/right)
                    tb_size = border_unit_converter(values[0]) or default_size
                    lr_size = border_unit_converter(values[1]) or default_size
                    for side in ("top", "bottom"):
                        borders[side]["size"] = tb_size
                    for side in ("left", "right"):
                        borders[side]["size"] = lr_size
                elif len(values) == 4:
                    # Four widths (top, right, bottom, left)
                    for side, val in zip(border_sides, values):
                        borders[side]["size"] = border_unit_converter(val) or default_size

            # Case 2: 'border-width', 'border-color', 'border-style' — apply to all sides
            elif css_prop in ("border-width", "border-color", "border-style"):
                prop_type = css_prop.split("-")[1]
                if prop_type == "width":
                    size = border_unit_converter(css_value) or default_size
                    for side in border_sides:
                        borders[side]["size"] = size
                elif prop_type == "color":
                    color = utils.parse_color(css_value, return_hex=True)
                    for side in border_sides:
                        borders[side]["color"] = color
                elif prop_type == "style":
                    style = parse_border_style(css_value.lower())
                    for side in border_sides:
                        borders[side]["style"] = style

            # Case 3: 'border-top', 'border-right-width', etc.
            else:
                parts = css_prop.split("-")
                side = parts[1]
                prop_type = parts[2] if len(parts) > 2 else None
                if prop_type == "width":
                    borders[side]["size"] = border_unit_converter(css_value) or default_size
                elif prop_type == "color":
                    borders[side]["color"] = utils.parse_color(css_value, return_hex=True)
                elif prop_type == "style":
                    borders[side]["style"] = parse_border_style(css_value.lower())
                else:
                    # Full side shorthand in any order
                    size, style, color = parse_border_value(css_value)
                    borders[side].update({"size": size, "style": style, "color": color})

        # Check if w:tcBorders exists, otherwise create it
        tcBorders = tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # Apply borders to the cell
        for side, border_info in borders.items():
            if border_info["size"] > 0:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), border_info["style"])  # Set border style
                border.set(qn("w:sz"), str(border_info["size"] * 8))  # Word uses eighths of a point
                border.set(qn("w:color"), border_info["color"].replace('#', ''))  # Set border color
                tcBorders.append(border)

    def add_bookmark(self, bookmark_name):
        """Adds a word bookmark to an existing paragraph"""
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(self.bookmark_id))
        bookmark_start.set(qn('w:name'), bookmark_name)

        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(self.bookmark_id))

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()

        self.paragraph._element.insert(0, bookmark_start)
        self.paragraph._element.append(bookmark_end)
        self.bookmark_id += 1

    def apply_styles_to_run(self, run, style, isCustom=False):
        if isCustom:
            try:
                run.style = style
                return
            except KeyError:
                print(f"Warning: Character style '{style}' not found in document.")
                return
            except ValueError as e:
                if "need type CHARACTER" in str(e):
                    print(
                        f"Warning: '{style}' is a paragraph style, not a character style."
                    )
                    print(
                        "For inline elements like <code>, please create a character style in Word."
                    )

        if not style or not hasattr(run, 'font'):
            return

        # Find current paragraph and run position
        if not hasattr(self, 'paragraph') or self.paragraph is None:
            return

        paragraph_id = id(self.paragraph)
        if paragraph_id not in self.paragraph_span_styles:
            self.paragraph_span_styles[paragraph_id] = {}

        # The current run is the last one in the paragraph
        run_index = len(self.paragraph.runs) - 1

        if run_index not in self.paragraph_span_styles[paragraph_id]:
            self.paragraph_span_styles[paragraph_id][run_index] = set()

        for style_name, style_value in style.items():
            if style_name in constants.RUN_STYLES:
                if style_name.startswith('background-color') and style_value in ('inherit', 'initial'):
                    continue

                self.paragraph_span_styles[paragraph_id][run_index].add(style_name)

                if style_name == 'text-decoration':
                    # If span sets text-decoration shorthand, it conflicts with all text-decoration-* properties
                    self.paragraph_span_styles[paragraph_id][run_index].add('text-decoration-line')
                    self.paragraph_span_styles[paragraph_id][run_index].add('text-decoration-style')
                    self.paragraph_span_styles[paragraph_id][run_index].add('text-decoration-color')
                elif style_name.startswith('text-decoration-'):
                    pass

        for style_name, style_value in style.items():
            if style_name in constants.PARAGRAPH_FORMAT_STYLES:
                continue
            elif style_name in constants.RUN_STYLES:
                handler = getattr(self, constants.RUN_STYLES[style_name])
                param_name = style_name.replace('-', '_')
                handler(run=run, **{param_name: style_value})
            else:
                logging.warning(f"Warning: Unrecognized style '{style_name}', will be skipped.")

    def apply_styles_to_paragraph(self, paragraph, style, isCustom=False):
        if isCustom:
            try:
                paragraph.style = style
                return
            except KeyError:
                print(f"Warning: Style '{style}' not found in document.  Using default.")
                return

        if not style or not hasattr(paragraph, 'paragraph_format'):
            return

        for style_name, style_value in style.items():
            if style_name in constants.PARAGRAPH_FORMAT_STYLES:
                handler = getattr(self, constants.PARAGRAPH_FORMAT_STYLES[style_name])
            elif style_name in constants.PARAGRAPH_RUN_STYLES:
                handler = getattr(self, constants.PARAGRAPH_RUN_STYLES[style_name])
            else:
                logging.warning(f"Warning: Unrecognized paragraph style '{style_name}', will be skipped.")
                continue

            handler(
                paragraph=paragraph,
                style_name=style_name,
                value=style_value,
                all_styles=style
            )

    def _apply_alignment_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        align = utils.remove_important_from_style(value)

        if 'center' in align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'left' in align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif 'right' in align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif 'justify' in align:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _apply_line_height_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        line_height = utils.remove_important_from_style(value)

        if line_height in ('normal', 'inherit'):
            paragraph.paragraph_format.line_spacing = None
        else:
            try:
                if line_height.replace('.', '').replace('%', '').isdigit():
                    multiplier = float(line_height[:-1]) / 100.0 if line_height.endswith('%') else float(line_height)
                    paragraph.paragraph_format.line_spacing = multiplier
                else:
                    converted = utils.unit_converter(line_height, target_unit="pt")
                    if converted is not None:
                        paragraph.paragraph_format.line_spacing = converted
            except (ValueError, TypeError):
                paragraph.paragraph_format.line_spacing = None

    def _apply_margins_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        style_name = kwargs['style_name']
        all_styles = kwargs['all_styles']

        margin_left = all_styles.get('margin-left')
        margin_right = all_styles.get('margin-right')

        if margin_left and margin_right:
            if 'auto' in margin_left and 'auto' in margin_right:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return

        if style_name == 'margin-left' and margin_left and 'auto' not in margin_left:
            paragraph.paragraph_format.left_indent = utils.unit_converter(margin_left)

        if style_name == 'margin-right' and margin_right and 'auto' not in margin_right:
            paragraph.paragraph_format.right_indent = utils.unit_converter(margin_right)

    def _apply_text_indent_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        indent_value = utils.remove_important_from_style(value)

        paragraph.paragraph_format.first_line_indent = utils.unit_converter(indent_value, target_unit="pt")

    def _apply_font_weight_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        font_weight = utils.remove_important_from_style(value).lower()

        paragraph_id = id(paragraph)
        paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

        for i, run in enumerate(paragraph.runs):
            if i in paragraph_spans and 'font-weight' in paragraph_spans[i]:
                continue

            self._apply_font_weight_to_run(
                run=run,
                font_weight=font_weight,
            )

    def _apply_font_weight_to_run(self, **kwargs):
        font_weight = kwargs['font_weight']
        run = kwargs['run']
        if font_weight in ('bold', 'bolder', '700', '800', '900'):
            run.font.bold = True
        elif font_weight in ('normal', 'lighter', '400', '300', '100'):
            run.font.bold = False
        # Note: Decide what to do for values between 400-700
        elif font_weight.isdigit():
            weight = int(font_weight)
            run.font.bold = weight >= 700

    def _apply_font_style_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        font_style = utils.remove_important_from_style(value).lower()

        paragraph_id = id(paragraph)
        paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

        for i, run in enumerate(paragraph.runs):
            if i in paragraph_spans and 'font-style' in paragraph_spans[i]:
                continue

            self._apply_font_style_to_run(
                run=run,
                font_style=font_style,
            )

    def _apply_font_style_to_run(self, **kwargs):
        font_style = kwargs['font_style']
        run = kwargs['run']

        if font_style in ('italic', 'oblique'):
            run.font.italic = True
        elif font_style == 'normal':
            run.font.italic = False

    def _apply_font_size_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        font_size = utils.remove_important_from_style(value).lower()

        if font_size in constants.FONT_SIZES_NAMED:
            font_size = constants.FONT_SIZES_NAMED[font_size]

        paragraph_id = id(paragraph)
        paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

        for i, run in enumerate(paragraph.runs):
            if i in paragraph_spans and 'font-size' in paragraph_spans[i]:
                continue

            self._apply_font_size_to_run(
                run=run,
                font_size=font_size,
            )

    def _apply_font_size_to_run(self, **kwargs):
        run = kwargs['run']
        font_size = kwargs['font_size']

        font_size = utils.remove_important_from_style(font_size).lower()
        font_size = utils.adapt_font_size(font_size)

        try:
            if font_size in ('normal', 'initial', 'inherit'):
                run.font.size = None
            else:
                converted_size = utils.unit_converter(font_size, target_unit="pt")
                if converted_size is not None:
                    run.font.size = converted_size

        except (ValueError, TypeError) as e:
            logging.warning(f"Warning: Could not parse font-size '{font_size}': {e}")

    def _apply_font_family_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        font_family = utils.remove_important_from_style(value).strip()

        paragraph_id = id(paragraph)
        paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

        for i, run in enumerate(paragraph.runs):
            if i in paragraph_spans and 'font-family' in paragraph_spans[i]:
                continue

            self._apply_font_family_to_run(
                run=run,
                font_family=font_family,
            )

    def _apply_font_family_to_run(self, **kwargs):
        run = kwargs['run']
        font_family = kwargs['font_family']

        if not font_family or font_family in ('inherit', 'initial', 'unset'):
            return

        try:
            font_families = [f.strip().strip('"\'') for f in font_family.split(',')]

            for font_name in font_families:
                if font_name and font_name not in ('inherit', 'initial', 'unset', 'serif', 'sans-serif', 'monospace',
                                                   'cursive', 'fantasy', 'system-ui'):
                    run.font.name = font_name
                    break
                elif font_name in ('serif', 'sans-serif', 'monospace'):
                    run.font.name = constants.GENERIC_FONT_STYLES[font_name]
                    break

        except (AttributeError, Exception) as e:
            logging.warning(f"Warning: Could not apply font-family '{font_family}': {e}")

    def _apply_color_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        all_styles = kwargs['all_styles']
        color_value = utils.remove_important_from_style(all_styles.get('color', '')).lower().strip()
        if color_value in ('inherit', 'initial', 'transparent', 'currentcolor'):
            return

        paragraph_id = id(paragraph)
        paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

        for i, run in enumerate(paragraph.runs):
            if i in paragraph_spans and 'color' in paragraph_spans[i]:
                continue
            self._apply_color_to_run(
                run=run,
                color=color_value,
            )

    def _apply_color_to_run(self, **kwargs):
        run = kwargs['run']
        color_value = kwargs['color']
        try:
            colors = utils.parse_color(color_value)
            run.font.color.rgb = RGBColor(*colors)
        except (ValueError, AttributeError) as e:
            logging.warning(f"Could not apply color '{color_value}': {e}")

    def _apply_text_transform_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        text_transform = utils.remove_important_from_style(value).lower()

        paragraph_id = id(paragraph)
        paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

        for i, run in enumerate(paragraph.runs):
            if i in paragraph_spans and 'text-transform' in paragraph_spans[i]:
                continue

            self._apply_text_transform_to_run(
                run=run,
                text_transform=text_transform,
            )

    def _apply_text_transform_to_run(self, **kwargs):
        run = kwargs['run']
        text_transform = kwargs['text_transform']

        if not run.text:
            return

        try:
            if text_transform == 'uppercase':
                run.text = run.text.upper()
            elif text_transform == 'lowercase':
                run.text = run.text.lower()
            elif text_transform == 'capitalize':
                run.text = run.text.title()
            elif text_transform in ('none', 'initial', 'inherit'):
                # No transformation needed
                pass
            elif text_transform in ('full-width', 'math-auto', 'full-size-kana'):
                logging.warning(f"Warning: Unsupported text transform '{text_transform}'")

        except (AttributeError, Exception) as e:
            logging.warning(f"Warning: Could not apply text-transform '{text_transform}': {e}")

    def _apply_text_decoration_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        all_styles = kwargs['all_styles']

        # Initialize decorations
        decorations = {
            'line_type': None,
            'line_style': None,
            'color': None
        }

        if 'text-decoration' in all_styles:
            text_decoration_value = utils.remove_important_from_style(all_styles['text-decoration']).lower()
            decorations = utils.parse_text_decoration(text_decoration_value)

        if 'text-decoration-line' in all_styles:
            line_value = utils.remove_important_from_style(all_styles['text-decoration-line']).lower()
            decorations['line_type'] = line_value

        if 'text-decoration-style' in all_styles:
            style_value = utils.remove_important_from_style(all_styles['text-decoration-style']).lower()
            decorations['line_style'] = style_value

        if 'text-decoration-color' in all_styles:
            color_value = utils.remove_important_from_style(all_styles['text-decoration-color']).lower()
            decorations['color'] = color_value

        paragraph_id = id(paragraph)
        paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

        for i, run in enumerate(paragraph.runs):
            span_styles = paragraph_spans.get(i, set())

            # If span has text-decoration shorthand, skip entirely
            if 'text-decoration' in span_styles:
                continue

            if decorations['line_type'] and 'text-decoration-line' not in span_styles:
                self._apply_text_decoration_line_to_run(
                    run=run,
                    text_decoration_line=decorations['line_type'],
                )

            if decorations['line_style'] and 'text-decoration-style' not in span_styles:
                self._apply_text_decoration_style_to_run(
                    run=run,
                    text_decoration_style=decorations['line_style'],
                )

            if decorations['color'] and 'text-decoration-color' not in span_styles:
                self._apply_text_decoration_color_to_run(
                    run=run,
                    text_decoration_color=decorations['color'],
                )

    def _apply_text_decoration_to_run(self, **kwargs):
        run = kwargs['run']
        text_decoration = kwargs['text_decoration']

        if not text_decoration:
            return

        decorations = utils.parse_text_decoration(text_decoration)
        if decorations['line_type']:
            self._apply_text_decoration_line_to_run(
                run=run,
                text_decoration_line=decorations['line_type'],
            )

        if decorations['line_style']:
            self._apply_text_decoration_style_to_run(
                run=run,
                text_decoration_style=decorations['line_style'],
            )

        if decorations['color']:
            self._apply_text_decoration_color_to_run(
                run=run,
                text_decoration_color=decorations['color'],
            )

    def _apply_text_decoration_line_to_run(self, **kwargs):
        run = kwargs['run']
        text_decoration_line = kwargs['text_decoration_line']

        if text_decoration_line in constants.FONT_UNDERLINE:
            if text_decoration_line == 'underline':
                run.font.underline = True
                run.font.strike = False
            elif text_decoration_line == 'line-through':
                run.font.strike = True
                run.font.underline = False
        elif text_decoration_line == 'none':
            run.font.underline = False
            run.font.strike = False
        else:
            logging.warning(f"Warning: Unsupported text decoration '{text_decoration_line}'")

    def _apply_text_decoration_style_to_run(self, **kwargs):
        run = kwargs['run']
        text_decoration_style = kwargs['text_decoration_style']
        if not text_decoration_style or run.font.underline is False:
            return False

        should_apply = False
        if run.font.underline:
            should_apply = True
        elif hasattr(self.paragraph, '_pending_styles'):
            for pending_style in self.paragraph._pending_styles:
                if 'text-decoration' in pending_style or 'text-decoration-line' in pending_style:
                    should_apply = True
                    break

        if not should_apply:
            return False
        try:
            run.font.underline = constants.FONT_UNDERLINE_STYLES[text_decoration_style]
        except KeyError:
            logging.warning(f"Warning: Style not recognized'{text_decoration_style}', defaulting to single line.")

        # Mark that we applied a text-decoration style by adding text-decoration-line to span_styles
        paragraph_id = id(self.paragraph)
        run_index = len(self.paragraph.runs) - 1
        if paragraph_id in self.paragraph_span_styles and run_index in self.paragraph_span_styles[paragraph_id]:
            self.paragraph_span_styles[paragraph_id][run_index].add('text-decoration-line')
        return True

    def _apply_text_decoration_color_to_run(self, **kwargs):
        run = kwargs['run']
        text_decoration_color = kwargs['text_decoration_color']

        if not text_decoration_color or not utils.is_color(text_decoration_color):
            return

        color_hex = utils.parse_color(text_decoration_color, return_hex=True)
        rPr = run._r.get_or_add_rPr()
        u = rPr.find(qn('w:u'))
        if u is not None:
            u.set(qn('w:color'), color_hex.upper().lstrip('#'))

    def _apply_background_color_paragraph(self, **kwargs):
        paragraph = kwargs['paragraph']
        value = kwargs['value']

        background_color = utils.remove_important_from_style(value).lower().strip()

        if background_color in ('inherit', 'initial'):
            return
        elif background_color in ('transparent', 'none'):
            logging.warning(f"Warning: Unsupported background color '{background_color}'")
            return

        try:
            color_hex = utils.parse_color(background_color, return_hex=True)
            if not color_hex:
                return

            paragraph_id = id(paragraph)
            paragraph_spans = self.paragraph_span_styles.get(paragraph_id, {})

            for i, run in enumerate(paragraph.runs):
                if i in paragraph_spans and 'background-color' in paragraph_spans[i]:
                    continue
                self._apply_background_color_to_run(
                    run=run,
                    background_color=background_color,
                )

        except Exception as e:
            logging.warning(f"Could not apply background-color to paragraph: {e}")

    def _apply_background_color_to_run(self, **kwargs):
        run = kwargs['run']
        background_color = kwargs['background_color']
        try:
            if background_color in ('inherit', 'initial'):
                return
            elif background_color in ('transparent', 'none'):
                logging.warning(f"Warning: Unsupported background color '{background_color}'")
                return

            color_hex = utils.parse_color(background_color, return_hex=True)
            if not color_hex:
                return

            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color_hex.lstrip('#'))

            r_pr = run._element.get_or_add_rPr()

            # Remove existing shading
            existing_shd = r_pr.find(qn('w:shd'))
            if existing_shd is not None:
                r_pr.remove(existing_shd)

            r_pr.append(shd)

        except Exception as e:
            logging.warning(f"Could not apply background-color to run: {e}")

    def add_text_align_or_margin_to(self, obj, style):
        """Styles that can be applied on multiple objects"""
        if 'text-align' in style:
            align = utils.remove_important_from_style(style['text-align'])

            if 'center' in align:
                obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif 'left' in align:
                obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif 'right' in align:
                obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif 'justify' in align:
                obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if 'margin-left' in style and 'margin-right' in style:
            if 'auto' in style['margin-left'] and 'auto' in style['margin-right']:
                obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif 'margin-left' in style:
            obj.left_indent = utils.unit_converter(style['margin-left'])

    def add_styles_to_table_cell(self, styles, doc_cell, cell_row):
        """Styles that must be applied specifically in a _Cell object"""
        # Set background color
        if 'background-color' in styles:
            self.set_cell_background(doc_cell, styles['background-color'])

        # Set width (approximate, since DOCX uses different units)
        if 'width' in styles:
            doc_cell.width = utils.unit_converter(styles['width'])

        # Set height (due word limitations, cannot set individually cell height, only whole row)
        if 'height' in styles:
            cell_row.height = utils.unit_converter(styles['height'])

        # Set text color
        if 'color' in styles:
            color = utils.parse_color(styles['color'])
            if color:
                for paragraph in doc_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = color

        # Set vertical align (for individual cells)
        if 'vertical-align' in styles:
            align = utils.remove_important_from_style(styles['vertical-align'])

            if 'top' in align:
                doc_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            elif 'middle' in align:
                doc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            elif 'bottom' in align:
                doc_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        # Set borders
        if any('border' in style for style in styles.keys()):
            self.set_cell_borders(doc_cell, styles)

        self.add_text_align_or_margin_to(doc_cell.paragraphs[0], styles)

    def add_styles_to_run(self, style):
        if 'font-size' in style:
            font_size = utils.remove_important_from_style(style['font-size'])
            # Adapt font_size when text, ex.: small, medium, etc.
            font_size = utils.adapt_font_size(font_size)

            if font_size:
                for run in self.paragraph.runs:
                    run.font.size = utils.unit_converter(font_size)

        if 'color' in style:
            colors = utils.parse_color(style['color'])
            self.run.font.color.rgb = RGBColor(*colors)

        if 'background-color' in style:
            # This should stay here for div.
            # Little trick to apply background-color to paragraph
            # because `self.run.font.highlight_color`
            # has a very limited amount of colors
            color = utils.parse_color(style['background-color'], return_hex=True)

            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color.lstrip('#'))

            # Make sure the paragraph styling element exists
            self.paragraph.paragraph_format.element.get_or_add_pPr()

            # Append the shading element
            self.paragraph.paragraph_format.element.pPr.append(shd)

    def handle_li(self):
        '''
            Handle li tags
            source: https://stackoverflow.com/a/78685353/17274446
        '''
        list_depth = len(self.tags['list']) or 1
        list_type = self.tags['list'][-1] if self.tags['list'] else 'ul'
        level = min(list_depth, 3)
        style_key = list_type if level <= 1 else f"{list_type}{level}"
        list_style = constants.STYLES.get(style_key, 'List Number' if list_type == 'ol' else 'List Bullet')

        self.paragraph = self.doc.add_paragraph(style=list_style)
        self.in_li = True

        if list_type == "ol":
            # Use your current_ol_num_id (generated on <ol> open) as key
            ol_id = self.current_ol_num_id or -1

            if ol_id not in self._list_num_ids:
                # First time using this <ol> → create a new numId

                style_obj = self.paragraph.style
                num_id_style = None

                if hasattr(style_obj._element.pPr, 'numPr'):
                    num_id_style = style_obj._element.pPr.numPr.numId.val

                if num_id_style is not None:
                    ct_numbering = self.doc.part.numbering_part.numbering_definitions._numbering
                    ct_num = ct_numbering.num_having_numId(num_id_style)
                    abstractNumId = ct_num.abstractNumId.val

                    # Add new numId linked to same abstractNumId
                    ct_num_new = ct_numbering.add_num(abstractNumId)
                    new_num_id = ct_num_new.numId

                    # Apply startOverride for level 0
                    lvl_override = ct_num_new.add_lvlOverride(0)
                    start_override = lvl_override._add_startOverride()
                    start_override.val = 1

                    # Cache this new numId
                    self._list_num_ids[ol_id] = new_num_id
            else:
                new_num_id = self._list_num_ids[ol_id]

            # Assign this numId to the paragraph
            pPr = self.paragraph._p.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')

            numId_elem = OxmlElement('w:numId')
            numId_elem.set(qn('w:val'), str(new_num_id))

            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(level - 1))

            numPr.append(ilvl)
            numPr.append(numId_elem)
            pPr.append(numPr)

    def add_image_to_cell(self, cell, image, width=None, height=None):
        # python-docx doesn't have method yet for adding images to table cells. For now we use this
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image, width, height)

    def handle_img(self, current_attrs):
        if not self.include_images:
            self.skip = True
            self.skip_tag = 'img'
            return

        if 'src' not in current_attrs:
            self.doc.add_paragraph("<image: no_src>")
            return

        src = current_attrs['src']

        # added image dimension, interpreting values as pixel only
        height = utils.unit_converter(current_attrs['height']) if 'height' in current_attrs else None
        width = utils.unit_converter(current_attrs['width']) if 'width' in current_attrs else None

        # fetch image
        image = utils.fetch_image_data(src)

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()

        self.run = self.paragraph.add_run()

        # add image to doc
        if image:
            try:
                if isinstance(self.doc, docx.document.Document):
                    self.run.add_picture(image, width, height)
                else:
                    self.add_image_to_cell(self.doc, image, width, height)
            except FileNotFoundError:
                image = None

        if not image:
            if utils.is_url(src):
                self.doc.add_paragraph("<image: %s>" % src)
            else:
                # avoid exposing filepaths in document
                self.doc.add_paragraph("<image: %s>" % utils.get_filename_from_url(src))

        '''
        #adding style
        For right-alignment: `'float: right;'`
        For center-alignment: `'display: block; margin-left: auto; margin-right: auto;'`
        Everything else would be Left aligned
        '''
        if 'style' in current_attrs:
            style = current_attrs['style']
            image_alignment = utils.get_image_alignment(style)
            last_paragraph = self.doc.paragraphs[-1]
            if image_alignment == utils.ImageAlignment.RIGHT:
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if image_alignment == utils.ImageAlignment.CENTER:
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def handle_table(self, current_attrs):
        """
        To handle nested tables, we will parse tables manually as follows:
        Get table soup
        Create docx table
        Iterate over soup and fill docx table with new instances of this parser
        Tell HTMLParser to ignore any tags until the corresponding closing table tag
        """
        table_soup = self.tables[self.table_no]
        rows, cols = self.get_table_dimensions(table_soup)
        self.table = self.doc.add_table(rows, cols)

        if self.table_style:
            # Available Table Styles
            # https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html#table-styles-in-default-template
            try:
                # Fixed 'style lookup by style_id is deprecated.'
                # https://stackoverflow.com/a/29567907/17274446
                self.table_style = ' '.join(re.findall(r'[A-Z][a-z]*|[0-9]', self.table_style))
                # Available Table Styles
                # https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html#table-styles-in-default-template
                self.table.style = self.table_style
            except KeyError as e:
                raise ValueError(f"Unable to apply style {self.table_style}.") from e

        # Reference:
        # https://python-docx.readthedocs.io/en/latest/api/table.html#cell-objects
        used_cells = [[False] * cols for _ in range(rows)]
        for cell_row, row in enumerate(self.get_table_rows(table_soup)):
            col_offset = 0  # Shift index if some columns are occupied
            for col in self.get_table_columns(row):
                while used_cells[cell_row][col_offset]:
                    col_offset += 1

                current_row = cell_row
                current_col = col_offset

                cell_html = self.get_cell_html(col)
                if col.name == 'th':
                    cell_html = "<b>%s</b>" % cell_html

                # Get _Cell object from table based on cell_row and cell_col
                docx_cell = self.table.cell(current_row, current_col)

                # Reference:
                # https://python-docx.readthedocs.io/en/latest/dev/analysis/features/table/cell-merge.html
                rowspan = int(col.get('rowspan', 1))
                colspan = int(col.get('colspan', 1))

                if rowspan > 1 or colspan > 1:
                    docx_cell = docx_cell.merge(
                        self.table.cell(
                            current_row + (rowspan - 1),
                            current_col + (colspan - 1)
                        )
                    )

                    # Mark all merged cells as used
                    for r in range(current_row, current_row + rowspan):
                        for c in range(current_col, current_col + colspan):
                            used_cells[r][c] = True
                else:
                    used_cells[current_row][current_col] = True

                # Parse cell styles
                cell_styles = utils.parse_dict_string(col.get('style', ''))

                if 'width' in cell_styles or 'height' in cell_styles:
                    self.table.autofit = False
                    self.table.allow_autofit = False

                child_parser = HtmlToDocx()
                child_parser.copy_settings_from(self)
                child_parser.add_html_to_cell(cell_html, docx_cell)
                child_parser.add_styles_to_table_cell(cell_styles, docx_cell, self.table.rows[cell_row])

                col_offset += colspan  # Move to the next real column

        if 'style' in current_attrs and self.table:
            style = utils.parse_dict_string(current_attrs['style'])
            self.add_text_align_or_margin_to(self.table, style)

        # skip all tags until corresponding closing tag
        self.instances_to_skip = len(table_soup.find_all('table'))
        self.skip_tag = 'table'
        self.skip = True
        self.table = None

    def handle_div(self, current_attrs):
        self.paragraph = self.doc.add_paragraph()

        # handle page break
        if 'style' in current_attrs and 'page-break-after: always' in current_attrs['style']:
            self.doc.add_page_break()

    def handle_hr(self):
        # This implementation was taken from:
        # https://github.com/python-openxml/python-docx/issues/105#issuecomment-62806373
        self.paragraph = self.doc.add_paragraph()
        pPr = self.paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(
            pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
            'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
            'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
            'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange'
        )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    def handle_link(self, href, text, tooltip=None):
        """
        A function that places a hyperlink within a paragraph object.

        Args:
            href: A string containing the required url.
            text: The text displayed for the url.
            tooltip: The text displayed when holder link.
        """
        is_external = href.startswith('http') if href else False
        hyperlink = OxmlElement('w:hyperlink')

        if is_external:
            # Create external hyperlink
            rel_id = self.paragraph.part.relate_to(
                href,
                docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                is_external=True
            )

            # Create the w:hyperlink tag and add needed values
            hyperlink.set(qn('r:id'), rel_id)
        else:
            # Create internal hyperlink (anchor)
            hyperlink.set(qn('w:anchor'), href.replace('#', ''))

        if tooltip is not None:
            # set tooltip to hyperlink
            hyperlink.set(qn('w:tooltip'), tooltip)

        # Create sub-run
        subrun = self.paragraph.add_run()
        rPr = OxmlElement('w:rPr')

        # add default color
        c = OxmlElement('w:color')
        c.set(qn('w:val'), "0000EE")
        rPr.append(c)

        # add underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        subrun._r.append(rPr)
        subrun._r.text = text

        # Add subrun to hyperlink
        hyperlink.append(subrun._r)

        # Add hyperlink to run
        self.paragraph._p.append(hyperlink)

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return
        if tag == 'head':
            self.skip = True
            self.skip_tag = tag
            self.instances_to_skip = 0
            return
        elif tag == 'body':
            return

        current_attrs = dict(attrs)

        if tag == 'span':
            # Parse inline styles if present to check for !important
            if "style" in current_attrs:
                normal_styles, important_styles = utils.parse_inline_styles(
                    current_attrs["style"]
                )
                # Store normal styles to apply to runs
                if normal_styles:
                    self.pending_inline_styles = normal_styles
                # Store important styles to apply after parent's processing
                if important_styles:
                    self.pending_important_styles = important_styles
            self.tags['span'].append(current_attrs)
            return
        elif tag in ['ol', 'ul']:
            if tag == 'ol':
                # Assign new ID if it's a fresh top-level list
                self.list_restart_counter += 1
                self.current_ol_num_id = self.list_restart_counter
            else:
                self.current_ol_num_id = None  # unordered list

            self.tags['list'].append(tag)
            return # don't apply styles for now
        elif tag == 'br':
            try:
                self.run.add_break()
            except AttributeError:
                self.paragraph = self.doc.add_paragraph()
                self.run = self.paragraph.add_run()
                self.run.add_break()
            return

        self.tags[tag] = current_attrs

        # Control custom_style based on the Options.  Default is True on both.
        custom_style = (
            self.get_word_style_for_element(tag, current_attrs)
            if (self.use_styles or self.use_tag_overrides)
            else None
        )

        if custom_style:
            valid_style = utils.check_style_exists(self.doc, custom_style)
            if not valid_style:
                custom_style = None

        if tag in ["p", "pre"]:
            if not self.in_li:
                self.paragraph = self.doc.add_paragraph()
                style_to_apply = self.pending_div_style or custom_style or self.default_paragraph_style
                if style_to_apply:
                    self.apply_styles_to_paragraph(self.paragraph, style_to_apply, True)
                # DON'T clear pending_div_style here - it should persist for all child paragraphs
                # It will be cleared when the div closes in handle_endtag

            # Parse inline styles on the paragraph itself to apply to runs within
            if "style" in current_attrs:
                normal_styles, important_styles = utils.parse_inline_styles(
                    current_attrs["style"]
                )
                if normal_styles:
                    self.pending_inline_styles = normal_styles
                if important_styles:
                    self.pending_important_styles = important_styles
        elif tag == "div":
            if custom_style and not self.in_li:
                self.pending_div_style = custom_style
            else:
                self.handle_div(current_attrs)
        elif tag == 'li':
            self.handle_li()
            if custom_style and self.paragraph:
                self.apply_styles_to_paragraph(self.paragraph, custom_style, True)

        elif tag == 'hr':
            self.handle_hr()

        elif re.match('h[1-9]', tag):
            if isinstance(self.doc, docx.document.Document):
                if custom_style:
                    self.paragraph = self.doc.add_paragraph()
                    self.apply_styles_to_paragraph(self.paragraph, custom_style, True)
                else:
                    h_size = int(tag[1])
                    self.paragraph = self.doc.add_heading(level=min(h_size, 9))
            else:
                self.paragraph = self.doc.add_paragraph()
                if custom_style:
                    self.apply_styles_to_paragraph(self.paragraph, custom_style, True)

        elif tag == 'img':
            self.handle_img(current_attrs)
            self.paragraph = self.doc.paragraphs[-1]

        elif tag == 'table':
            if self.include_tables:
                self.handle_table(current_attrs)
                return

        elif tag == "code":
            if custom_style:
                self.pending_character_style = custom_style
            if "style" in current_attrs:
                normal_styles, important_styles = utils.parse_inline_styles(
                    current_attrs["style"]
                )
                if normal_styles:
                    self.pending_inline_styles = normal_styles
                if important_styles:
                    self.pending_important_styles = important_styles
            return
        # Commented this out.  Line breaks are already handled by try/except blocks (see the br code above).  In addition to this, I'm fixing the tests added by the previous release.
        # # set new run reference point in case of leading line breaks
        # if tag in ["p", "li", "pre"]:
        #     self.run = self.paragraph.add_run()

        if 'id' in current_attrs:
            self.add_bookmark(current_attrs['id'])

        # add style
        if not self.include_styles:
            return

        if 'style' in current_attrs and self.paragraph and (tag in ['p'] or re.match(r'h[1-9]', tag)):
            if not hasattr(self.paragraph, '_pending_styles'):
                self.paragraph._pending_styles = []
            style = utils.parse_dict_string(current_attrs['style'])
            self.paragraph._pending_styles.append(style)
        elif 'style' in current_attrs and self.paragraph:
            style = utils.parse_dict_string(current_attrs['style'])
            self.add_text_align_or_margin_to(self.paragraph.paragraph_format, style)

    def handle_endtag(self, tag):
        # Clear pending character style and inline styles when closing inline elements
        if tag == "code":
            self.pending_character_style = None
            self.pending_inline_styles = None
            self.pending_important_styles = None

        # Clear important styles when closing span
        if tag == "span":
            self.pending_important_styles = None

        # Clear pending div style when closing a div
        if tag == "div":
            self.pending_div_style = None

        # Clear pending inline styles when closing paragraph elements
        if tag in ["p", "pre"]:
            self.pending_inline_styles = None
            self.pending_important_styles = None

        if re.match('h[1-9]', tag):
            self.pending_inline_styles = None
            self.pending_important_styles = None

        if self.skip:
            if not tag == self.skip_tag:
                return

            if self.instances_to_skip > 0:
                self.instances_to_skip -= 1
                return

            self.skip = False
            self.skip_tag = None
            self.paragraph = None

        if tag == 'span':
            if self.tags['span']:
                self.tags['span'].pop()
                return
        elif tag in ['ol', 'ul']:
            utils.remove_last_occurence(self.tags['list'], tag)
            if tag == 'ol':
                self._list_num_ids.pop(self.current_ol_num_id, None)
                self.current_ol_num_id = None
            return
        elif tag == 'table':
            if self.include_tables:
                self.table_no += 1
                self.table = None
                self.paragraph = None
        elif tag == 'li':
            self.in_li = False

        if tag in ['p', 'pre'] or re.match(r'h[1-9]', tag):
            if hasattr(self.paragraph, '_pending_styles'):
                for style in self.paragraph._pending_styles:
                    self.apply_styles_to_paragraph(self.paragraph, style)
                # Clear the pending styles
                del self.paragraph._pending_styles
            self.paragraph_span_styles.clear()

        if tag in self.tags:
            self.tags.pop(tag)
        # maybe set relevant reference to None?

    def handle_data(self, data):
        if self.skip:
            return

        # Only remove white space if we're not in a pre block.
        if 'pre' not in self.tags:
            # remove leading and trailing whitespace in all instances
            data = utils.remove_whitespace(data, True, True)

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()

        # There can only be one nested link in a valid html document
        # You cannot have interactive content in an A tag, this includes links
        # https://html.spec.whatwg.org/#interactive-content
        link = self.tags.get('a', {})
        href = link.get('href', None)
        title = link.get('title', None)

        if link and href:
            self.handle_link(href, data, title)
            return

        # If there's a link, dont put the data directly in the run
        self.run = self.paragraph.add_run(data)

        for span in self.tags['span']:
            if 'style' in span:
                span_style = utils.parse_dict_string(span['style'])
                self.apply_styles_to_run(self.run, span_style)

                for tag, attrs in self.tags.items():
                    if tag == 'div' and 'style' in attrs:
                        div_style = utils.parse_dict_string(attrs['style'])

                        for span_style_name in span_style.keys():
                            if span_style_name in div_style:
                                del div_style[span_style_name]

                        self.tags[tag]['style'] = utils.dict_to_style_string(div_style)

        for tag, attrs in self.tags.items():
            if tag in constants.FONT_STYLES:
                font_style = constants.FONT_STYLES[tag]
                setattr(self.run.font, font_style, True)

            if tag in constants.FONT_NAMES:
                font_name = constants.FONT_NAMES[tag]
                self.run.font.name = font_name

            if 'style' in attrs and (tag in ['div', 'li', 'pre']):
                style = utils.parse_dict_string(attrs['style'])
                self.add_styles_to_run(style)

    def handle_comment(self, data):
        """
        Render HTML comments (<!-- text -->) as visible green text,
        starting with "#", similar to how code editors show comments.
        """
        if not self.include_html_comments:
            return

        # Put comment on its own line
        self.paragraph = self.doc.add_paragraph()

        comment_text = f"# {data.strip()}"
        run = self.paragraph.add_run(comment_text)

        # Style: Green color to mimic HTML comment styling
        dark_ish_green = "#008000"
        run.font.color.rgb = utils.parse_color(dark_ish_green)
        run.italic = True  # makes it feel more like a comment

    def ignore_nested_tables(self, tables_soup):
        """
        Returns array containing only the highest level tables
        Operates on the assumption that bs4 returns child elements immediately after
        the parent element in `find_all`. If this changes in the future, this method will need to be updated

        :return:
        """
        new_tables = []
        nest = 0
        for table in tables_soup:
            if nest:
                nest -= 1
                continue
            new_tables.append(table)
            nest = len(table.find_all('table'))
        return new_tables

    def get_table_rows(self, table_soup):
        # If there's a header, body, footer or direct child tr tags, add row dimensions from there
        return table_soup.select(', '.join(self.table_row_selectors), recursive=False)

    def get_table_columns(self, row):
        # Get all columns for the specified row tag.
        return row.find_all(['th', 'td'], recursive=False) if row else []

    def get_table_dimensions(self, table_soup):
        # Get rows for the table
        rows = self.get_table_rows(table_soup)
        # Table is either empty or has non-direct children between table and tr tags
        # Thus the row dimensions and column dimensions are assumed to be 0
        # A table can have a varying number of columns per row,
        # so it is important to find the maximum number of columns in any row
        if not rows:
            return 0, 0

        default_span = 1
        max_cols = 0
        max_rows = len(rows)

        for row_idx, row in enumerate(rows):
            cols = self.get_table_columns(row)
            # Handle colspan
            row_col_count = sum(int(col.get('colspan', default_span)) for col in cols)
            max_cols = max(max_cols, row_col_count)

            # Handle rowspan
            for col in cols:
                rowspan = int(col.get('rowspan', default_span))
                if rowspan > default_span:
                    max_rows = max(max_rows, row_idx + rowspan)

        return max_rows, max_cols

    def get_tables(self) -> None:
        if not hasattr(self, 'soup'):
            self.options['tables'] = False
            return

        self.tables = self.ignore_nested_tables(self.soup.find_all('table'))
        self.table_no = 0

    def run_process(self, html: str) -> None:
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')

            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_document(self, html: str, document) -> None:
        if not isinstance(html, str):
            raise ValueError(f'First argument needs to be a {str}')
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError(f'Second argument needs to be a {docx.document.Document}')

        self.set_initial_attrs(document)
        self.run_process(html)

    def add_html_to_cell(self, html: str, cell: docx.table._Cell) -> None:
        if not isinstance(cell, docx.table._Cell):
            raise ValueError(f'Second argument needs to be a {docx.table._Cell}')

        unwanted_paragraph = cell.paragraphs[0]
        utils.delete_paragraph(unwanted_paragraph)
        self.set_initial_attrs(cell)
        self.run_process(html)
        # cells must end with a paragraph or will get message about corrupt file
        # https://stackoverflow.com/a/29287121
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')

    def parse_html_file(self, filename_html: str, filename_docx, encoding: str = 'utf-8') -> None:
        with open(filename_html, 'r', encoding=encoding) as infile:
            html = infile.read()

        self.set_initial_attrs()
        self.run_process(html)

        if not filename_docx:
            path, filename = os.path.split(filename_html)
            filename_docx = f'{path}/new_docx_file_{filename}'

        self.save(filename_docx)

    def parse_html_string(self, html: str) -> docx.document.Document:
        self.set_initial_attrs()
        self.run_process(html)
        return self.doc

if __name__ == '__main__':
    arg_parser = argparse.ArgumentParser(description='Convert .html file into .docx file with formatting')
    arg_parser.add_argument('filename_html', help='The .html file to be parsed')
    arg_parser.add_argument(
        'filename_docx',
        nargs='?',
        help='The name of the .docx file to be saved. Default new_docx_file_[filename_html]',
        default=None
    )
    arg_parser.add_argument('--bs', action='store_true',
                            help='Attempt to fix html before parsing. Requires bs4. Default True')

    args = vars(arg_parser.parse_args())
    file_html = args.pop('filename_html')
    html_parser = HtmlToDocx()
    html_parser.parse_html_file(file_html, **args)
